import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_table_from_md(doc, table_lines):
    if not table_lines: return
    rows = []
    for line in table_lines:
        if line.strip().startswith('|--'): continue
        cells = [c.strip() for c in line.split('|') if c.strip() or line.count('|') > 1]
        if cells: rows.append(cells)
    if not rows: return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            if j < len(table.columns):
                cell = table.cell(i, j)
                cell.text = cell_data
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs: run.bold = True

def process_text_formatting(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2]); run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1]); run.italic = True
        else:
            subparts = re.split(r'(\$.*?\$)', part)
            for subpart in subparts:
                if subpart.startswith('$') and subpart.endswith('$'):
                    run = paragraph.add_run(subpart[1:-1]); run.italic = True
                    run.font.color.rgb = RGBColor(0, 0, 139)
                else:
                    paragraph.add_run(subpart)

def convert_md_to_docx_advanced(md_path, docx_path):
    doc = Document()
    style = doc.styles['Normal']; font = style.font; font.name = 'Segoe UI'; font.size = Pt(10.5)
    with open(md_path, 'r', encoding='utf-8') as f: content = f.read()
    lines = content.replace('\r\n', '\n').split('\n')
    in_code_block = False; title_found = False; i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('```'):
            in_code_block = not in_code_block; i += 1; continue
        if in_code_block:
            p = doc.add_paragraph(line); p.style = 'No Spacing'
            run = p.runs[0] if p.runs else p.add_run(line); run.font.name = 'Consolas'; run.font.size = Pt(9)
            i += 1; continue
        if '|' in line and (i + 1 < len(lines) and '|--' in lines[i+1]):
            table_lines = [line]; i += 1
            while i < len(lines) and '|' in lines[i]: table_lines.append(lines[i]); i += 1
            add_table_from_md(doc, table_lines); continue
        if line.strip() == '$$':
            i += 1; math_lines = []
            while i < len(lines) and line.strip() != '$$':
                line = lines[i]
                if line.strip() == '$$': break
                math_lines.append(line); i += 1
            p = doc.add_paragraph('\n'.join(math_lines)); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0] if p.runs else p.add_run('\n'.join(math_lines)); run.italic = True; run.font.color.rgb = RGBColor(0, 0, 139)
            i += 1; continue
        if line.startswith('# '):
            heading = line[2:].strip()
            if not title_found:
                p = doc.add_heading(heading, 0); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; title_found = True
            else: doc.add_heading(heading, level=1)
        elif line.startswith('## '): doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '): doc.add_heading(line[4:].strip(), level=3)
        elif line.strip().startswith(('- ', '* ')):
            p = doc.add_paragraph(style='List Bullet'); process_text_formatting(p, line.strip()[2:])
        elif re.match(r'^\s*\d+\. ', line.strip()):
            p = doc.add_paragraph(style='List Number'); process_text_formatting(p, re.sub(r'^\s*\d+\. ', '', line.strip()))
        elif line.strip():
            p = doc.add_paragraph(); process_text_formatting(p, line.strip())
        i += 1
    doc.save(docx_path)
    print(f"Updated: {docx_path}")

def main():
    # Update README and PAPER as well
    convert_md_to_docx_advanced('README.md', 'docs/README.docx')
    convert_md_to_docx_advanced('PARALLAX.md', 'docs/PAPER_v1.8.5.docx')
    for f in os.listdir('docs'):
        if f.endswith('.md'):
            docx_name = f.replace('.md', '.docx')
            if "NOVEL_CONTRIBUTIONS" in docx_name:
                print(f"Skipping locked file: {docx_name}")
                continue
            try:
                convert_md_to_docx_advanced(os.path.join('docs', f), os.path.join('docs', docx_name))
            except Exception as e:
                print(f"Failed to update {docx_name}: {e}")

if __name__ == "__main__":
    main()
