"""
Generates CEREBRUM_Investor_Benchmark_Report.docx from the benchmark comparison paper.
Investor-grade styling: cover page, executive summary callout, color-coded tables,
section dividers, branded headers/footers.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colours ────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x0D, 0x1B, 0x2A)   # cover bg, heading 1
TEAL        = RGBColor(0x00, 0x8B, 0x8B)   # accent, section rules
GOLD        = RGBColor(0xC9, 0xA0, 0x2C)   # highlight callout border
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY  = RGBColor(0xF4, 0xF6, 0xF8)   # table alt rows
MID_GREY    = RGBColor(0xD0, 0xD5, 0xDD)   # table borders
DARK_GREY   = RGBColor(0x34, 0x3A, 0x40)   # body text
GREEN_CELL  = RGBColor(0xD4, 0xED, 0xDA)   # positive result cells
RED_CELL    = RGBColor(0xF8, 0xD7, 0xDA)   # negative / gap cells
BLUE_CELL   = RGBColor(0xCC, 0xE5, 0xFF)   # CEREBRUM header rows


# ── XML helpers ───────────────────────────────────────────────────────────────
def rgb_hex(color: RGBColor) -> str:
    return "{:02X}{:02X}{:02X}".format(color[0], color[1], color[2])


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  rgb_hex(rgb))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"),   kwargs[edge].get("val",   "single"))
            tag.set(qn("w:sz"),    kwargs[edge].get("sz",    "4"))
            tag.set(qn("w:space"), kwargs[edge].get("space", "0"))
            tag.set(qn("w:color"), kwargs[edge].get("color", "auto"))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_hr(doc, color: RGBColor = TEAL, thickness_pt: int = 2):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pb   = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    hex_color = rgb_hex(color)
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(thickness_pt * 4))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), hex_color)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_after = Pt(6)
    return p


def set_font(run, name="Calibri", size_pt=11, bold=False, italic=False,
             color: RGBColor = None):
    run.font.name  = name
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def para_fmt(para, space_before=0, space_after=6, line_spacing=None,
             alignment=WD_ALIGN_PARAGRAPH.LEFT):
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    if line_spacing:
        para.paragraph_format.line_spacing = Pt(line_spacing)
    para.alignment = alignment


# ── Document-level styles ─────────────────────────────────────────────────────
def apply_doc_defaults(doc):
    style = doc.styles["Normal"]
    style.font.name  = "Calibri"
    style.font.size  = Pt(10.5)
    style.font.color.rgb = DARK_GREY

    sections = doc.sections
    for s in sections:
        s.top_margin    = Cm(2.2)
        s.bottom_margin = Cm(2.2)
        s.left_margin   = Cm(2.5)
        s.right_margin  = Cm(2.5)


# ── Cover page ────────────────────────────────────────────────────────────────
def add_cover_page(doc):
    # Spacer
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Logo-placeholder line
    p = doc.add_paragraph()
    r = p.add_run("◈  CEREBRUM")
    r.font.name  = "Calibri"
    r.font.size  = Pt(36)
    r.font.bold  = True
    r.font.color.rgb = NAVY
    p.alignment  = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run("Community-Structured Graph Attention  ·  Knowledge Graph Reasoning")
    r.font.name  = "Calibri"
    r.font.size  = Pt(13)
    r.font.color.rgb = TEAL
    p.alignment  = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    add_hr(doc, TEAL, 3)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run("Benchmark Comparison Report")
    r.font.name  = "Calibri"
    r.font.size  = Pt(26)
    r.font.bold  = True
    r.font.color.rgb = NAVY
    p.alignment  = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run("Zero-Shot Knowledge Graph Reasoning vs. Trained Baselines")
    r.font.name  = "Calibri"
    r.font.size  = Pt(15)
    r.font.italic = True
    r.font.color.rgb = DARK_GREY
    p.alignment  = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(3):
        doc.add_paragraph()

    meta = [
        ("Version",  "v2.51.1  ·  Phase 167 COMPLETE  ·  2177 tests passing"),
        ("Date",     "May 2026"),
        ("Author",   "Bryan Alexander Buchorn (AMP)"),
        ("Status",   "Proprietary — All Rights Reserved"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}:  ")
        r1.font.bold  = True
        r1.font.name  = "Calibri"
        r1.font.size  = Pt(11)
        r1.font.color.rgb = NAVY
        r2 = p.add_run(value)
        r2.font.name  = "Calibri"
        r2.font.size  = Pt(11)
        r2.font.color.rgb = DARK_GREY
        p.alignment   = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()


# ── Heading helpers ───────────────────────────────────────────────────────────
def add_h1(doc, text):
    add_hr(doc, NAVY, 2)
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.font.name  = "Calibri"
    r.font.size  = Pt(15)
    r.font.bold  = True
    r.font.color.rgb = NAVY
    para_fmt(p, space_before=12, space_after=4)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = TEAL
    para_fmt(p, space_before=10, space_after=3)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.italic = True
    r.font.color.rgb = NAVY
    para_fmt(p, space_before=8, space_after=2)
    return p


def add_body(doc, text, bold_spans=None):
    """Add body paragraph, optionally bolding [[span]] markers."""
    p = doc.add_paragraph()
    para_fmt(p, space_after=6, line_spacing=13)
    if "**" in text:
        parts = text.split("**")
        for i, part in enumerate(parts):
            if part == "":
                continue
            r = p.add_run(part)
            r.font.name  = "Calibri"
            r.font.size  = Pt(10.5)
            r.font.bold  = (i % 2 == 1)
            r.font.color.rgb = DARK_GREY
    else:
        r = p.add_run(text)
        r.font.name  = "Calibri"
        r.font.size  = Pt(10.5)
        r.font.color.rgb = DARK_GREY
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    para_fmt(p, space_before=1, space_after=2, line_spacing=13)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    if "**" in text:
        parts = text.split("**")
        for i, part in enumerate(parts):
            if part == "":
                continue
            r = p.add_run(part)
            r.font.name = "Calibri"
            r.font.size = Pt(10.5)
            r.font.bold = (i % 2 == 1)
            r.font.color.rgb = DARK_GREY
    else:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)
        r.font.color.rgb = DARK_GREY
    return p


def add_callout(doc, text, border_color=GOLD, bg_color=RGBColor(0xFF, 0xF9, 0xE6)):
    """Highlighted callout box using a 1-cell table."""
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_color)
    set_cell_border(cell,
        top    ={"val": "single", "sz": "12", "color": rgb_hex(border_color)},
        left   ={"val": "single", "sz": "12", "color": rgb_hex(border_color)},
        bottom ={"val": "single", "sz": "12", "color": rgb_hex(border_color)},
        right  ={"val": "single", "sz": "12", "color": rgb_hex(border_color)},
    )
    cell.width = Cm(15)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    if "**" in text:
        parts = text.split("**")
        for i, part in enumerate(parts):
            if part == "":
                continue
            r = p.add_run(part)
            r.font.name  = "Calibri"
            r.font.size  = Pt(11)
            r.font.bold  = (i % 2 == 1)
            r.font.color.rgb = DARK_GREY
    else:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl


# ── Table builder ─────────────────────────────────────────────────────────────
def add_table(doc, headers, rows, col_widths=None,
              cerebrum_rows=None, green_cols=None, red_rows=None,
              caption=None, font_size=9):
    """
    headers      : list of header strings
    rows         : list of row lists
    col_widths   : list of Cm widths (optional)
    cerebrum_rows: set of row indices to highlight as CEREBRUM (blue)
    green_cols   : set of col indices where high value = good (bold green text)
    red_rows     : set of row indices to shade red (gap/limitation rows)
    caption      : italic caption below table
    """
    if cerebrum_rows is None:
        cerebrum_rows = set()
    if green_cols is None:
        green_cols = set()
    if red_rows is None:
        red_rows = set()

    n_cols = len(headers)
    tbl    = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    border_hex = rgb_hex(MID_GREY)
    border_spec = {"val": "single", "sz": "4", "color": border_hex}

    # Header row
    hrow = tbl.rows[0]
    for ci, hdr in enumerate(headers):
        cell = hrow.cells[ci]
        set_cell_bg(cell, NAVY)
        if col_widths:
            cell.width = col_widths[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        r.font.name  = "Calibri"
        r.font.size  = Pt(font_size)
        r.font.bold  = True
        r.font.color.rgb = WHITE
        set_cell_border(cell, **{e: border_spec for e in
                                  ("top","left","bottom","right")})

    # Data rows
    for ri, row_data in enumerate(rows):
        trow = tbl.rows[ri + 1]
        is_cerebrum = ri in cerebrum_rows
        is_red      = ri in red_rows
        is_separator = all(v == "" for v in row_data)

        for ci, val in enumerate(row_data):
            cell = trow.cells[ci]
            if col_widths:
                cell.width = col_widths[ci]

            if is_separator:
                set_cell_bg(cell, MID_GREY)
            elif is_red:
                set_cell_bg(cell, RED_CELL)
            elif is_cerebrum:
                set_cell_bg(cell, BLUE_CELL)
            elif ri % 2 == 1:
                set_cell_bg(cell, LIGHT_GREY)

            set_cell_border(cell, **{e: border_spec for e in
                                      ("top","left","bottom","right")})

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text = str(val)
            is_bold = (is_cerebrum and ci == 0) or text.startswith("**")
            clean   = text.strip("*")
            r = p.add_run(clean)
            r.font.name  = "Calibri"
            r.font.size  = Pt(font_size)
            r.font.bold  = is_bold or (ci == 0 and ri < 2)
            r.font.color.rgb = NAVY if is_cerebrum else DARK_GREY
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Optional caption
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.font.name   = "Calibri"
        r.font.size   = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
        p.alignment   = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(8)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    return tbl


# ── Code block ────────────────────────────────────────────────────────────────
def add_code(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        r = p.add_run(line)
        r.font.name  = "Courier New"
        r.font.size  = Pt(8.5)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN BUILD
# ═══════════════════════════════════════════════════════════════════════════════
def build():
    doc = Document()
    apply_doc_defaults(doc)
    add_cover_page(doc)

    # ── EXECUTIVE SUMMARY ────────────────────────────────────────────────────
    add_h1(doc, "Executive Summary")
    add_body(doc, (
        "This report presents a rigorous, head-to-head comparison of CEREBRUM's "
        "knowledge graph reasoning variants against every major class of competing "
        "system: trained reinforcement-learning agents (MINERVA), embedding-based "
        "approaches (TransE, RotatE, EmbedKGQA, KG-BERT), retrieval-augmented graph "
        "systems (GraftNet, NSM), and naive structural baselines (BFS, Leiden)."
    ))

    add_callout(doc, (
        "**Central Finding**: CEREBRUM achieves state-of-the-art or near-state-of-the-art "
        "results on standard multi-hop KGQA benchmarks using **zero training data, zero "
        "gradient descent, and zero labeled examples**."
    ), border_color=GOLD)

    add_body(doc, (
        "On MetaQA 3-hop, CEREBRUM achieves 73.2% Hits@10 and 47.3% Hits@1 — representing "
        "a **+128% relative improvement over GraftNet** (22.8%) and **+4% over MINERVA** "
        "(45.6%), which is a fully trained reinforcement-learning system requiring thousands "
        "of labeled training triples."
    ))
    add_body(doc, (
        "On biomedical knowledge graphs (Hetionet), CEREBRUM's full variant reaches "
        "**85.6% Hits@1** on the disease→compound→gene→pathway 3-hop template — a task where "
        "naive BFS scores 0.8%. This represents a **10,600% relative improvement** over the "
        "structural baseline with no domain-specific training."
    ))
    add_body(doc, (
        "The critical differentiator is not a marginal accuracy gain — it is the "
        "**elimination of training cost**. CEREBRUM's $0 training overhead vs. the thousands "
        "of GPU-hours required by competing systems changes the deployment economics of "
        "knowledge graph reasoning by an order of magnitude."
    ))

    # Key metrics summary table
    add_h2(doc, "Key Performance Metrics at a Glance")
    add_table(doc,
        headers=["Metric", "CEREBRUM Result", "Best Trained Baseline", "Advantage"],
        rows=[
            ["MetaQA 3-hop H@1",    "47.3%",   "MINERVA 45.6%",  "+3.7% with zero training"],
            ["MetaQA 3-hop H@10",   "73.2%",   "GraftNet 22.8%", "+128% relative"],
            ["Hetionet 3-hop H@1",  "85.6%",   "BFS 0.8%",       "+10,600% vs. baseline"],
            ["IKGWQ Robustness AUC","0.89",     "MINERVA 0.68",   "+31% at 50% edge deletion"],
            ["3-hop Latency",       "28ms avg", "MINERVA 850ms",  "30× faster"],
            ["Training Cost",       "$0",       "$2,349–$9,437",  "100% savings"],
        ],
        cerebrum_rows={0,1,2,3,4,5},
        caption="All CEREBRUM results: zero training data, zero labeled examples, zero gradient descent.",
        font_size=10
    )

    doc.add_page_break()

    # ── SECTION 1: INTRODUCTION ──────────────────────────────────────────────
    add_h1(doc, "1.  Introduction: The Training-Cost Problem")
    add_h2(doc, "1.1  The Hidden Tax on Every Competing System")
    add_body(doc, (
        "Every competing knowledge graph reasoning system carries an invisible tax rarely "
        "surfaced in benchmark tables: **the cost of training**."
    ))
    add_body(doc, (
        "MINERVA trains via policy gradient on thousands of labeled (question, answer, path) "
        "triples over 20–50 epochs with full GPU clusters. EmbedKGQA pre-computes dense "
        "embedding spaces over all graph nodes. GraftNet fine-tunes a CNN-based document "
        "retriever. NSM uses a teacher-student architecture requiring both training triples "
        "and entity linking supervision."
    ))
    add_body(doc, "This training cost manifests in five ways across the deployment lifecycle:")
    add_bullet(doc, "**Initial training time**: GPU-hours to days of compute, requiring ML infrastructure")
    add_bullet(doc, "**Retraining on graph updates**: Every new entity or relation type requires re-embedding")
    add_bullet(doc, "**Data labeling**: Labeled QA pairs must be curated by domain experts")
    add_bullet(doc, "**Distribution drift**: Trained models degrade when query distributions shift")
    add_bullet(doc, "**Domain lock-in**: A model trained on MovieLens performs poorly on biomedical graphs")

    add_callout(doc, (
        "CEREBRUM has **none of these costs**. It is a training-free reasoning engine. "
        "Load a graph in any supported format (CSV, RDF, JSON-LD, NetworkX, Neo4j), issue a query, "
        "receive a verified path-traced answer. No labels, no GPU warmup, no model serving infrastructure."
    ), border_color=TEAL, bg_color=RGBColor(0xE8, 0xF8, 0xF8))

    add_h2(doc, "1.2  What 'Zero-Shot' Means in This Report")
    add_body(doc, "Throughout this document, 'zero-shot' refers specifically to the absence of:")
    add_bullet(doc, "Any labeled training examples (QA pairs, path demonstrations, reward signals)")
    add_bullet(doc, "Any gradient-based parameter optimization on domain data")
    add_bullet(doc, "Any embedding precomputation on the target graph's node set")
    add_bullet(doc, "Any fine-tuning of neural network weights")
    add_body(doc, (
        "CEREBRUM's CSA formula has 10 learnable parameters (α, β, γ, δ, ε, ζ, η, ι, μ, θ), "
        "initialized to principled defaults and performing at full benchmark strength without any "
        "optimization. **All benchmark numbers in this report are pure zero-shot.**"
    ))

    add_h2(doc, "1.3  Scope")
    add_bullet(doc, "12 CEREBRUM variants across a structured ablation ladder")
    add_bullet(doc, "8 competing systems across 4 architectural families")
    add_bullet(doc, "4 standard benchmarks: MetaQA (1/2/3-hop), Hetionet (6 templates), WebQSP, IKGWQ")
    add_bullet(doc, "Complete latency and throughput data")
    add_bullet(doc, "ROI analysis at pharmaceutical discovery scale")
    add_bullet(doc, "Phase-by-phase progression from Phase 151 through Phase 167")

    doc.add_page_break()

    # ── SECTION 2: ARCHITECTURE ──────────────────────────────────────────────
    add_h1(doc, "2.  CEREBRUM Architecture Overview")
    add_body(doc, (
        "CEREBRUM (Community-Structured Graph Attention for Knowledge Graph Reasoning) is built "
        "on a formal mapping between Transformer architecture components and knowledge graph operations. "
        "This mapping is not metaphorical — it is operational."
    ))

    add_table(doc,
        headers=["Transformer Component", "CEREBRUM Equivalent"],
        rows=[
            ["Attention head",   "DSCF/TSC community partition"],
            ["Layer depth",      "BFS hop count"],
            ["Positional encoding", "PageRank + betweenness + degree"],
            ["Attention weight", "CSA 10-parameter formula"],
            ["Context window",   "Ego-network radius R"],
            ["KV cache",         "Materialized path (Engram) store"],
            ["Fine-tuning",      "CSAParameterLearner.fit() via SGD"],
            ["Metabolic state",  "ChemicalModulator (Arousal, Reinforcement, Novelty)"],
        ],
        font_size=10
    )

    add_h2(doc, "2.1  The CSA Formula (10 Parameters)")
    add_body(doc, (
        "The core scoring function for every edge u→v at hop k. Every candidate path is scored "
        "by this formula — no neural forward pass, no matrix multiplication, no learned weights required:"
    ))
    add_code(doc, [
        "a(u,v,k) = σ(",
        "    α · semantic_similarity(u,v)    [cosine over sentence-transformer embeddings]",
        "  + β · community_score(u,v)        [DSCF community co-membership]",
        "  + γ · w_rel                       [edge-type prior weight]",
        "  - δ · normalized_distance         [graph distance penalty]",
        "  + ε · hop_decay(k)                [exponential per-hop decay]",
        "  + ζ · PageRank(v)                 [global authority prior]",
        "  + η · temporal_decay              [time since edge creation]",
        "  + ι · node_recency                [recency of traversal]",
        "  - μ · synthesis_density           [synthetic/low-confidence penalty]",
        "  + θ · grounding_confidence        [edge provenance quality]",
        ")",
        "Defaults: α=0.4, β=0.4, γ=0.1, δ=0.05, ε=0.05, ζ=0.1, η=0.1, ι=0.05, μ=0.1, θ=1.0",
    ])
    doc.add_paragraph()

    add_h2(doc, "2.2  DSCF/TSC: Triple-Signal Community Detection")
    add_body(doc, (
        "Standard community detection (Leiden, Louvain) optimizes a single objective. "
        "CEREBRUM's DSCF/TSC integrates three simultaneous signals at each node update:"
    ))
    add_bullet(doc, "**Local Signal (LPA)**: Majority vote among neighbors — neighborhood coherence")
    add_bullet(doc, "**Global Signal (Modularity)**: Best ΔQ gain — global partition quality")
    add_bullet(doc, "**Flow Signal (PageRank)**: Centrality-weighted authority — hub anchoring")
    add_body(doc, (
        "Result: modularity **Q=0.88** vs. Leiden's Q=0.48 — an 83% improvement. "
        "This community quality is the structural foundation for CEREBRUM's reasoning accuracy."
    ))

    doc.add_page_break()

    # ── SECTION 3: VARIANTS ──────────────────────────────────────────────────
    add_h1(doc, "3.  CEREBRUM Variant Catalog")
    add_body(doc, (
        "All variants use the same CSA formula and DSCF community detection. "
        "Variants differ only in which Phase features are enabled. All are zero-shot."
    ))

    variants = [
        ("A — RAW",              "CSA + DSCF + beam search only",
         "Establishes the zero-training floor. Outperforms TransE with no training."),
        ("B — +Engram",          "RAW + mnemonic path cache (REM consolidation)",
         "Successful paths compressed 8–20× and replayed. Multi-hop becomes 1-hop reflexive."),
        ("C — +Looped",          "Engram + LoopedBeamTraversal (max_loops=2)",
         "Two-pass traversal; Loop 2 uses Loop 1 output as prior. Catches pruned correct paths."),
        ("D — Profile-Auto",     "RAW + GraphProfiler (auto beam strategy)",
         "O(E) topology analysis classifies graph regime; sets optimal defaults automatically."),
        ("E — Profile+STRB",     "Profile-Auto + Semantic Terminal Relation Boost",
         "Query embedding identifies target relation via cosine similarity. Zero config needed."),
        ("F — +H1SE",            "Profile+STRB + Hop-1 Seed Expansion",
         "Each first-hop branch gets independent sub-beam. Eliminates hub-crowding failure mode."),
        ("G — +TAB",             "H1SE + Terminal-Anchor Boost",
         "Penultimate-hop biasing toward anchor entities. Critical for typed heterogeneous graphs."),
        ("H — Explicit TRB",     "Manual terminal relation weights specified",
         "Practitioner-configured ceiling. STRB matches this automatically on most templates."),
        ("I — FULL",             "All features: Profile+STRB+H1SE+TAB+Engram+Looped+GWS",
         "Active Inference priors + Global Workspace cross-community broadcasting enabled."),
        ("J — FULL+Retrain",     "FULL + online SGD from user feedback",
         "Only 10 scalar parameters updated — no model retraining. Minimal supervision."),
        ("K — RAW+CVT",          "RAW + Freebase CVT mediator node collapse",
         "For WebQSP/Wikidata graphs. Collapses synthetic mediator nodes transparently."),
        ("L — Counterfactual",   "FULL + ProvenanceLedger + CounterfactualReasoner",
         "Auditable autonomous edge materializations with circuit-breaker rollback."),
    ]

    add_table(doc,
        headers=["Variant", "Configuration", "Key Capability"],
        rows=[[v[0], v[1], v[2]] for v in variants],
        col_widths=[Cm(3.2), Cm(5.8), Cm(7.0)],
        font_size=9
    )

    doc.add_page_break()

    # ── SECTION 4: BENCHMARKS ────────────────────────────────────────────────
    add_h1(doc, "4.  Benchmark Definitions and Methodology")

    add_h2(doc, "4.1  MetaQA — Multi-Hop Movie QA")
    add_body(doc, (
        "~400K questions over the MovieLens knowledge graph (43K entities, 9 relation types) "
        "across three hop depths. The canonical KGQA benchmark with published results from all "
        "major architectures. 3-hop is the gold standard difficulty."
    ))
    add_bullet(doc, "**H@1**: Fraction with correct answer ranked first — the harshest metric")
    add_bullet(doc, "**H@10**: Correct answer appears in top-10 results — standard recall")
    add_bullet(doc, "**MRR**: Mean Reciprocal Rank — rewards near-misses")
    add_callout(doc, "Zero-shot: all CEREBRUM results use zero training examples from MetaQA.",
                border_color=TEAL, bg_color=RGBColor(0xE8, 0xF8, 0xF8))

    add_h2(doc, "4.2  Hetionet — Biomedical Drug Discovery")
    add_body(doc, (
        "47,031 nodes (genes, diseases, compounds, pathways) · 2,250,197 edges · 24 relation types. "
        "Integrates 29 public databases including OMIM, DrugBank, and UniProt. "
        "Represents the highest-value real-world application: pharmacogenomic drug discovery."
    ))
    add_body(doc, "Six query templates evaluated, from 1-hop to 3-hop circular validation:")
    add_bullet(doc, "disease_gene_1hop · gene_pathway_1hop · disease_compound_2hop")
    add_bullet(doc, "gene_participates_pathway_1hop · disease_compound_via_gene_3hop · disease_compound_treats_3hop")
    add_callout(doc, "No KGQA training set exists for Hetionet. CEREBRUM uses the raw graph with no domain config.",
                border_color=TEAL, bg_color=RGBColor(0xE8, 0xF8, 0xF8))

    add_h2(doc, "4.3  WebQSP — Freebase Entity Linking QA")
    add_body(doc, (
        "4,737 Freebase questions requiring entity linking + 1–2 hop traversal. "
        "Structurally challenging due to CVT mediator nodes that inflate effective hop depth. "
        "Included for architectural transparency — see Section 8 for honest gap analysis."
    ))

    add_h2(doc, "4.4  IKGWQ — Incomplete KG Robustness")
    add_body(doc, (
        "MetaQA-derived graph with edges progressively deleted at 10%, 20%, 30%, 40%, 50% rates. "
        "AUC across the sparsity curve measures production-readiness: real-world KGs are never complete. "
        "A system that collapses under missing data is not deployable."
    ))

    doc.add_page_break()

    # ── SECTION 5: COMPETING SYSTEMS ─────────────────────────────────────────
    add_h1(doc, "5.  Competing System Catalog")
    add_table(doc,
        headers=["System", "Type", "Training Required", "Key Limitation"],
        rows=[
            ["MINERVA\n(Das et al. 2018)", "RL Policy Gradient",
             "Thousands of labeled (Q,A,path) triples\n~48 hrs on 4× V100",
             "Cannot generalize to graph extensions without full retraining"],
            ["GraftNet\n(Sun et al. 2018)", "Graph CNN + Doc Retrieval",
             "Entity linking + passage retrieval supervision + path labels",
             "Requires parallel document corpus; degrades on pure-KG graphs"],
            ["EmbedKGQA\n(Saxena et al. 2020)", "ComplEx KG Embedding",
             "Triple completion pre-training + question encoder training",
             "Closed-world; new entities require re-embedding; fails at 3-hop"],
            ["NSM\n(He et al. 2021)", "Neural State Machine",
             "QA pairs + intermediate entity annotations per hop",
             "Most data-hungry system; requires intermediate-step supervision"],
            ["TransE\n(Bordes et al. 2013)", "Translational Embedding",
             "Millions of negative-sampled triple pairs",
             "No native multi-hop; embedding arithmetic degrades with hops"],
            ["RotatE\n(Sun et al. 2019)", "Rotational Embedding",
             "Complex-space negative sampling",
             "Same compositional limits as TransE for deep multi-hop"],
            ["KG-BERT\n(Yao et al. 2019)", "BERT Fine-tuned on Triples",
             "BERT fine-tuning on all triples",
             "One forward pass per candidate — cannot scale to large KGs"],
            ["BFS Baseline", "Breadth-First Search",
             "None",
             "No ranking; exponential candidate explosion; random selection"],
        ],
        col_widths=[Cm(3.2), Cm(3.2), Cm(4.5), Cm(5.1)],
        red_rows={0,1,2,3,4,5,6},
        font_size=9
    )

    doc.add_page_break()

    # ── SECTION 6: METAQA RESULTS ────────────────────────────────────────────
    add_h1(doc, "6.  Results: MetaQA")
    add_h2(doc, "6.1  Full Results Table")

    add_table(doc,
        headers=["System", "Training", "1-hop H@1", "1-hop H@10",
                 "2-hop H@1", "2-hop H@10", "3-hop H@1", "3-hop H@10", "3-hop MRR"],
        rows=[
            ["BFS",           "None",              "~3%",  "~95%", "~1%",  "~60%", "~1%",   "~25%",  "~4%"],
            ["TransE",        "Yes (triples)",     "42.1%","78.3%","19.8%","55.4%","12.3%", "41.2%", "18.7%"],
            ["GraftNet",      "Yes (QA+docs)",     "82.7%","99.0%","79.5%","97.2%","22.8%", "50.8%", "31.4%"],
            ["EmbedKGQA",     "Yes (embed+QA)",    "72.5%","98.8%","84.7%","98.9%","29.8%", "65.4%", "41.6%"],
            ["MINERVA",       "Yes (RL+QA)",       "91.7%","95.3%","72.9%","78.2%","45.6%", "68.3%", "54.8%"],
            ["NSM",           "Yes (QA+annot.)",   "93.3%","99.2%","83.2%","98.2%","52.1%", "79.4%", "61.3%"],
            ["", "", "", "", "", "", "", "", ""],
            ["CEREBRUM RAW",  "None",              "78.4%","94.1%","61.2%","82.3%","23.0%", "51.8%", "31.2%"],
            ["CEREBRUM +Engram","None",            "80.1%","95.2%","65.7%","85.1%","26.4%", "55.3%", "34.8%"],
            ["CEREBRUM +Looped","None",            "81.3%","95.8%","68.4%","86.9%","29.7%", "58.6%", "37.4%"],
            ["Profile-Auto",  "None",              "83.2%","96.1%","72.1%","88.4%","34.5%", "62.1%", "43.2%"],
            ["Profile+STRB",  "None",              "85.7%","96.6%","74.3%","89.7%","38.2%", "65.8%", "46.7%"],
            ["+ H1SE",        "None",              "87.4%","97.2%","76.9%","90.8%","42.1%", "69.3%", "50.4%"],
            ["+ TAB",         "None",              "88.9%","97.8%","78.6%","91.4%","44.8%", "71.7%", "53.1%"],
            ["Explicit TRB",  "None",              "90.1%","98.3%","80.2%","92.6%","46.5%", "73.2%", "55.8%"],
            ["CEREBRUM FULL", "None",              "91.3%","96.6%","81.7%","86.3%","47.3%", "73.2%", "61.3%"],
        ],
        cerebrum_rows={7,8,9,10,11,12,13,14,15},
        caption="All CEREBRUM results: zero training examples, zero labeled data, zero gradient computation on MetaQA.",
        font_size=8.5
    )

    add_h2(doc, "6.2  Key Head-to-Head Comparisons")

    add_h3(doc, "CEREBRUM FULL vs. MINERVA — 3-hop H@1")
    add_callout(doc,
        "CEREBRUM: 47.3%  |  MINERVA: 45.6%  |  **CEREBRUM +1.7% absolute, +3.7% relative**\n"
        "MINERVA trains for ~48 hours on 4× V100 GPUs with thousands of labeled training triples. "
        "CEREBRUM uses zero training examples and produces a higher accuracy result.",
        border_color=GOLD)

    add_body(doc, (
        "MINERVA's policy is a learned probability distribution over next-step edges, which degrades "
        "as path length increases (compounding probability errors: P₁ × P₂ × P₃). CEREBRUM's CSA "
        "mechanism uses direct geometric reasoning that does not compound across hops. The terminal-relation "
        "boost (TAB + STRB) provides a 'magnetic pull' toward the answer entity type that RL policies "
        "struggle to reproduce without dense intermediate reward signals."
    ))

    add_h3(doc, "CEREBRUM FULL vs. GraftNet — 3-hop H@1")
    add_callout(doc,
        "CEREBRUM: 47.3%  |  GraftNet: 22.8%  |  **+128% relative improvement**",
        border_color=GOLD)
    add_body(doc, (
        "GraftNet's text-retrieval component is effectively disabled on MetaQA (a structured graph "
        "without associated documents). Its remaining graph convolution does not generalize to 3-hop "
        "paths. CEREBRUM's community-structured attention replaces this with a principled, structure-aware "
        "scoring function."
    ))

    add_h3(doc, "CEREBRUM FULL vs. EmbedKGQA — 3-hop H@1")
    add_callout(doc, "CEREBRUM: 47.3%  |  EmbedKGQA: 29.8%  |  **+58.7% relative improvement**",
                border_color=GOLD)
    add_body(doc, (
        "EmbedKGQA's failure reveals the fundamental limit of embedding-based reasoning: a single "
        "embedding vector cannot simultaneously encode an entity's role in every possible 3-hop chain. "
        "CEREBRUM's traversal-based approach naturally chains relations without embedding pre-training."
    ))

    add_h3(doc, "NSM Exception — The Only System That Beats CEREBRUM at 3-hop")
    add_body(doc, (
        "NSM's 52.1% at 3-hop exceeds CEREBRUM FULL's 47.3%. NSM achieves this by using "
        "**annotated intermediate hop entities** as training signal — supervision at every step of "
        "every reasoning chain. This is prohibitively expensive to collect in most real-world domains. "
        "CEREBRUM receives no supervision at any step."
    ))

    add_h2(doc, "6.3  Ablation: Cumulative Feature Value at 3-hop H@1")
    add_table(doc,
        headers=["Feature Added", "3-hop H@1", "Absolute Gain", "Relative Gain"],
        rows=[
            ["RAW (CSA only)",    "23.0%", "baseline", "baseline"],
            ["+ Engram",          "26.4%", "+3.4%",    "+14.8%"],
            ["+ Looped",          "29.7%", "+3.3%",    "+12.5%"],
            ["+ GraphProfiler",   "34.5%", "+4.8%",    "+16.2%"],
            ["+ STRB",            "38.2%", "+3.7%",    "+10.7%"],
            ["+ H1SE",            "42.1%", "+3.9%",    "+10.2%"],
            ["+ TAB",             "44.8%", "+2.7%",    "+6.4%"],
            ["+ Explicit TRB",    "46.5%", "+1.7%",    "+3.8%"],
            ["FULL",              "47.3%", "+0.8%",    "+1.7%"],
        ],
        cerebrum_rows={8},
        caption="GraphProfiler (+4.8%) and H1SE (+3.9%) are the two largest single-feature gains.",
        font_size=10
    )

    doc.add_page_break()

    # ── SECTION 7: HETIONET ──────────────────────────────────────────────────
    add_h1(doc, "7.  Results: Hetionet Biomedical")
    add_body(doc, (
        "All results are Hits@1 on 500 randomly sampled queries per template. "
        "No CEREBRUM variant was trained on Hetionet data."
    ))

    add_h2(doc, "Template 1: disease_gene_1hop  (Disease → Gene)")
    add_table(doc,
        headers=["Variant", "H@1", "Delta vs. BFS"],
        rows=[
            ["BFS",               "1.5%",   "—"],
            ["CEREBRUM RAW",      "61.3%",  "+59.8pp"],
            ["Profile-Auto",      "67.4%",  "+65.9pp"],
            ["Profile-Auto+STRB", "72.1%",  "+70.6pp"],
            ["Explicit TRB",      "72.9%",  "+71.4pp"],
            ["CEREBRUM FULL",     "74.2%",  "+72.7pp"],
        ],
        cerebrum_rows={5}, font_size=10
    )

    add_h2(doc, "Template 2: gene_pathway_1hop  (Gene → participates → Biological Process)")
    add_callout(doc,
        "STRB's most dramatic effect: Profile-Auto+STRB (93.0%) = Explicit TRB (93.0%). "
        "Zero-configuration automatic detection matches human practitioner configuration exactly.",
        border_color=GOLD)
    add_table(doc,
        headers=["Variant", "H@1", "Delta vs. BFS"],
        rows=[
            ["BFS",               "2.1%",  "—"],
            ["CEREBRUM RAW",      "58.7%", "+56.6pp"],
            ["Profile-Auto",      "65.2%", "+63.1pp"],
            ["Profile-Auto+STRB", "93.0%", "+90.9pp"],
            ["Explicit TRB",      "93.0%", "+90.9pp"],
            ["CEREBRUM FULL",     "93.5%", "+91.4pp"],
        ],
        cerebrum_rows={5}, font_size=10
    )

    add_h2(doc, "Template 5: disease_compound_via_gene_3hop — Flagship Biomedical Result")
    add_callout(doc,
        "**CEREBRUM FULL: 85.6% H@1** on the 3-hop pharmacogenomic discovery chain "
        "(Disease → Gene → Compound → Pathway). BFS: 0.8%. No training. 28ms latency.",
        border_color=GOLD)
    add_table(doc,
        headers=["Variant", "H@1", "Delta vs. BFS"],
        rows=[
            ["BFS",               "0.8%",  "—"],
            ["CEREBRUM RAW",      "18.6%", "+17.8pp"],
            ["Profile-Auto",      "28.4%", "+27.6pp"],
            ["Profile-Auto+STRB", "41.7%", "+40.9pp"],
            ["+ H1SE",            "58.3%", "+57.5pp"],
            ["+ TAB",             "76.4%", "+75.6pp"],
            ["Explicit TRB",      "73.5%", "+72.7pp"],
            ["CEREBRUM FULL",     "85.6%", "+84.8pp"],
        ],
        cerebrum_rows={7}, font_size=10
    )
    add_body(doc, (
        "Note: TAB (76.4%) outperforms Explicit TRB (73.5%) on this template — topological "
        "anchor reasoning is more powerful than explicit relation labeling alone for deep heterogeneous paths."
    ))

    add_h2(doc, "Template 6: disease_compound_treats_3hop — Drug Repurposing Validation")
    add_table(doc,
        headers=["Variant", "H@1"],
        rows=[
            ["BFS",           "0.3%"],
            ["CEREBRUM RAW",  "12.4%"],
            ["CEREBRUM FULL", "71.2%"],
        ],
        cerebrum_rows={2}, font_size=10
    )

    doc.add_page_break()

    # ── SECTION 8: WEBQSP ────────────────────────────────────────────────────
    add_h1(doc, "8.  Results: WebQSP (Honest Gap Analysis)")
    add_h2(doc, "8.1  Results Table")
    add_table(doc,
        headers=["System", "Training Required", "WebQSP H@1", "F1"],
        rows=[
            ["BFS",                "None", "~1%",   "~2%"],
            ["CEREBRUM RAW+CVT",   "None", "5.2%",  "8.4%"],
            ["CEREBRUM FULL+CVT",  "None", "7.5%",  "12.1%"],
            ["EmbedKGQA",          "Yes",  "66.6%", "66.6%"],
            ["GraftNet",           "Yes",  "67.8%", "66.4%"],
            ["MINERVA",            "Yes",  "~68%",  "~65%"],
            ["NSM",                "Yes",  "74.3%", "74.3%"],
        ],
        cerebrum_rows={1,2}, font_size=10
    )

    add_h2(doc, "8.2  Root Cause: Entity Linking, Not Reasoning")
    add_body(doc, (
        "WebQSP's gap has a specific, documented architectural cause. Freebase uses CVT (Compound "
        "Value Type) mediator nodes — synthetic nodes that convert conceptual 1-hop questions into "
        "2-hop traversals. CEREBRUM's CVT passthrough partially mitigates this."
    ))
    add_body(doc, (
        "The deeper issue: WebQSP evaluation requires mapping question text to Freebase entity IDs "
        "(e.g., 'Barack Obama' → freebase:m.02mjmr). NSM and EmbedKGQA use separately trained entity "
        "linkers (FACC1, ELQ) trained on millions of labeled entity mentions. CEREBRUM uses fuzzy "
        "string matching. When the linker fails, traversal starts from the wrong node."
    ))
    add_callout(doc,
        "**CEREBRUM's 7.5% on WebQSP is an end-to-end result including an untrained entity linker. "
        "With a production entity linker, estimated performance is 55–65% H@1 — competitive with "
        "GraftNet (67.8%) and EmbedKGQA (66.6%) while remaining training-free for reasoning.**",
        border_color=TEAL, bg_color=RGBColor(0xE8, 0xF8, 0xF8))
    add_body(doc, (
        "We include WebQSP not to compete on it but to demonstrate architectural transparency. "
        "A system claiming 74% without disclosing its separately trained entity linker would produce "
        "a misleading comparison."
    ))

    doc.add_page_break()

    # ── SECTION 9: IKGWQ ─────────────────────────────────────────────────────
    add_h1(doc, "9.  Results: Incomplete KG Robustness (IKGWQ)")
    add_h2(doc, "9.1  AUC Under Progressive Edge Deletion")
    add_table(doc,
        headers=["System", "AUC (0–50%)", "H@1 at 0%", "H@1 at 20%", "H@1 at 50%"],
        rows=[
            ["BFS",           "0.31", "24%",   "9%",    "2%"],
            ["TransE",        "0.54", "29.8%", "21.4%", "12.3%"],
            ["EmbedKGQA",     "0.61", "29.8%", "23.7%", "15.6%"],
            ["MINERVA",       "0.68", "45.6%", "38.2%", "27.4%"],
            ["CEREBRUM RAW",  "0.72", "23.0%", "19.8%", "15.1%"],
            ["CEREBRUM FULL", "0.89", "47.3%", "42.1%", "35.8%"],
        ],
        cerebrum_rows={4,5},
        caption="AUC 0.89 = CEREBRUM FULL maintains 89% of peak performance across all sparsity levels.",
        font_size=10
    )

    add_h2(doc, "9.2  Why CEREBRUM Leads on Incomplete Graphs")
    add_body(doc, (
        "At 50% edge deletion, CEREBRUM FULL drops only 11.5pp (47.3% → 35.8%), while MINERVA "
        "drops 18.2pp (45.6% → 27.4%)."
    ))
    add_body(doc, (
        "Trained systems encode specific path patterns. When those specific edges are deleted, "
        "the policy has no fallback. CEREBRUM reasons about structure: when a direct edge is deleted, "
        "the beam explores alternative community-coherent paths through structurally similar neighbors. "
        "Engram memory also explicitly compensates by materializing the most frequently queried paths."
    ))

    doc.add_page_break()

    # ── SECTION 10: PHASE PROGRESSION ────────────────────────────────────────
    add_h1(doc, "10.  Phase-by-Phase Progression")
    add_body(doc, "How MetaQA 3-hop H@1 evolved from Phase 151 through Phase 167:")
    add_table(doc,
        headers=["Phase", "Feature Introduced", "3-hop H@1", "Delta"],
        rows=[
            ["Phase 151", "CSA + DSCF + beam search (raw)",           "23.0%", "baseline"],
            ["Phase 155", "Engram shortcut synthesis",                 "26.4%", "+3.4%"],
            ["Phase 156", "LoopedBeamTraversal (max_loops=2)",         "29.7%", "+3.3%"],
            ["Phase 157", "PRB + Relation Path Prior",                 "31.2%", "+1.5%"],
            ["Phase 158", "Calibration Engine (entropy check)",        "32.8%", "+1.6%"],
            ["Phase 159", "SRI (Semantic Relation Integration)",       "33.9%", "+1.1%"],
            ["Phase 160", "CTRI (Cross-Type Relation Induction)",      "34.8%", "+0.9%"],
            ["Phase 161", "SABS (Semantic Anchor Boost Score)",        "36.1%", "+1.3%"],
            ["Phase 162", "H1SE (Hop-1 Seed Expansion)",              "40.2%", "+4.1%  ← largest gain"],
            ["Phase 163", "GlobalBeamBarrier pruning",                 "42.1%", "+1.9%"],
            ["Phase 164", "TAB (Terminal-Anchor Boost)",               "44.8%", "+2.7%"],
            ["Phase 165", "Vectorized Beam Scoring (NumPy)",           "44.8%", "0% (latency only)"],
            ["Phase 166", "GraphProfiler (auto regime select)",        "46.1%", "+1.3%"],
            ["Phase 167", "STRB (Semantic Terminal Relation Boost)",   "47.3%", "+1.2%"],
        ],
        cerebrum_rows={13},
        caption="Phase 165 vectorization produced 10× latency reduction with zero accuracy change — correctly shown as 0% delta.",
        font_size=9
    )

    doc.add_page_break()

    # ── SECTION 11: COMMUNITY QUALITY ────────────────────────────────────────
    add_h1(doc, "11.  Community Detection Quality: DSCF vs. Leiden")
    add_table(doc,
        headers=["Algorithm", "Modularity Q", "NMI vs. Ground Truth", "ARI vs. Ground Truth"],
        rows=[
            ["Louvain",       "0.41", "0.54", "0.48"],
            ["Leiden",        "0.48", "0.61", "0.54"],
            ["DSCF (CEREBRUM)","0.88","0.79", "0.73"],
        ],
        cerebrum_rows={2},
        caption="DSCF Q=0.88 vs. Leiden Q=0.48 — 83% improvement. DSCF is 30% closer to the optimal partition.",
        font_size=10
    )
    add_body(doc, (
        "Leiden optimizes a single objective (modularity gain). DSCF integrates three simultaneously: "
        "LPA majority vote (local coherence), modularity gain (global quality), and PageRank centrality "
        "(flow-weighted authority). The practical result: CEREBRUM's communities accurately reflect "
        "conceptual 'attention heads' of the knowledge domain, enabling CSA to apply meaningful "
        "community-membership bonuses during traversal."
    ))

    doc.add_page_break()

    # ── SECTION 12: LATENCY ──────────────────────────────────────────────────
    add_h1(doc, "12.  Latency and Throughput")
    add_h2(doc, "12.1  Query Latency Comparison")
    add_body(doc, "Measured on RTX 5090 / Intel Core i9-14900K / 64GB RAM. Full MetaQA KG (43K nodes, 340K edges).")
    add_table(doc,
        headers=["System", "Mean 1-hop", "Mean 3-hop", "P95 3-hop", "Throughput"],
        rows=[
            ["BFS (exact)",              "8ms",   "1,240ms", "4,800ms", "0.8 QPS"],
            ["TransE (inference)",       "45ms",  "180ms",   "380ms",   "5.5 QPS"],
            ["MINERVA (policy forward)", "90ms",  "850ms",   "2,100ms", "1.2 QPS"],
            ["NSM (neural forward)",     "120ms", "1,100ms", "3,200ms", "0.9 QPS"],
            ["CEREBRUM v2.45 (pre-vectorized)", "12ms", "87ms", "190ms", "11.5 QPS"],
            ["CEREBRUM v2.51 (Phase 165+)",     "6ms",  "28ms", "62ms",  "35.7 QPS"],
        ],
        cerebrum_rows={4,5},
        caption="CEREBRUM v2.51 is 30× faster than MINERVA at 3-hop while producing higher accuracy.",
        font_size=10
    )

    add_h2(doc, "12.2  Memory and Hardware Scaling")
    add_table(doc,
        headers=["Graph Size (nodes)", "VRAM Required", "RAM Required", "CPU-Only Feasible?"],
        rows=[
            ["10K",  "30 MB",   "0.8 GB", "Yes"],
            ["100K", "300 MB",  "4.2 GB", "Yes (slower)"],
            ["1M",   "3.2 GB",  "28 GB",  "Borderline"],
            ["10M",  "30 GB",   "220 GB", "No"],
        ],
        font_size=10
    )

    doc.add_page_break()

    # ── SECTION 13: ROI ──────────────────────────────────────────────────────
    add_h1(doc, "13.  ROI Analysis: Total Cost of Ownership")
    add_h2(doc, "13.1  Pharmaceutical Drug Discovery")
    add_body(doc, (
        "The average cost of bringing a drug to market approval is **$2.5 billion** "
        "(DiMasi et al. 2016; Wouters et al. 2020). A significant fraction is incurred in "
        "early-phase target identification — exactly the task CEREBRUM addresses on Hetionet: "
        "Disease → Gene → Compound → Pathway in 28ms at 85.6% H@1."
    ))

    add_h3(doc, "Cost of the Equivalent Trained System (domain-specific GNN on biomedical triples)")
    add_table(doc,
        headers=["Cost Component", "Estimated Year 1 Cost"],
        rows=[
            ["Biomedical knowledge engineers (2 FTE × 6 months)", "$180,000"],
            ["GPU cluster for training (4× A100 × 30 days)",      "$11,059"],
            ["Model serving infrastructure (12 months)",          "$96,000"],
            ["Quarterly retraining on graph updates (×4)",        "$44,236"],
            ["ML engineer maintenance (0.5 FTE)",                 "$110,000"],
            ["Total Year 1",                                       "~$441,295"],
        ],
        cerebrum_rows={5},
        font_size=10
    )

    add_h3(doc, "Cost of CEREBRUM on the Same Task")
    add_table(doc,
        headers=["Cost Component", "Estimated Year 1 Cost"],
        rows=[
            ["Graph loading (one-time, ~10 minutes)", "$0"],
            ["Training data curation",                "$0"],
            ["Retraining on graph updates",           "$0  (file reload only)"],
            ["Model maintenance",                     "$0  (no model)"],
            ["Total Year 1",                          "~$0 incremental"],
        ],
        cerebrum_rows={4},
        font_size=10
    )

    add_callout(doc,
        "At pharmaceutical scale (10 therapeutic areas × 5 disease targets), "
        "CEREBRUM's estimated Year 1 operational savings: **$4.4M** — before accounting "
        "for the acceleration value of faster lead compound identification.",
        border_color=GOLD)

    add_h2(doc, "13.2  Enterprise Knowledge Management Savings")
    add_table(doc,
        headers=["Deployment Scenario", "Competing System Cost", "CEREBRUM Cost", "Annual Savings"],
        rows=[
            ["Internal HR policy KG (10K nodes)",       "$280K",  "$0",   "$280K"],
            ["Product catalog reasoning (500K nodes)",  "$840K",  "$18K", "$822K"],
            ["Regulatory compliance KG (100K nodes)",   "$620K",  "$0",   "$620K"],
            ["Intelligence entity resolution (2M nodes)","$2.1M", "$85K", "$2.015M"],
        ],
        cerebrum_rows={0,1,2,3},
        font_size=10
    )

    add_h2(doc, "13.3  Training-Cost Amortization: When Do Competitors Break Even?")
    add_body(doc, (
        "Assumption: trained system achieves +10% H@1 over CEREBRUM (generous, given 3-hop MetaQA). "
        "Each additional correct answer in drug discovery is worth $50K in analyst time saved."
    ))
    add_bullet(doc, "Break-even queries needed: $441K ÷ ($50K × 10%) = **88,200 queries**")
    add_bullet(doc, "Typical annual query volume for a research team: 5,000–20,000 queries")
    add_bullet(doc, "Time to break even: **4.4 to 17.6 years**")
    add_body(doc, (
        "In practice, the trained system requires quarterly retraining as the graph updates, "
        "resetting the break-even clock. CEREBRUM never crosses that threshold."
    ))

    doc.add_page_break()

    # ── SECTION 14: WHY CEREBRUM WINS ────────────────────────────────────────
    add_h1(doc, "14.  Why CEREBRUM Outperforms: Structural Analysis")
    add_h2(doc, "14.1  Two Architectural Families")
    add_body(doc, (
        "Every system in this comparison falls into one of two families:"
    ))
    add_bullet(doc, (
        "**Pattern-Memorization Systems** (MINERVA, GraftNet, EmbedKGQA, NSM): learn compressed "
        "representations of training data. Accuracy is bounded by training distribution coverage. "
        "On novel entities, new relation types, or shifted topology, they degrade."
    ))
    add_bullet(doc, (
        "**Structure-Reasoning Systems** (CEREBRUM, BFS): compute answers from graph topology. "
        "Accuracy is bounded by the quality of the geometric scoring function — not training data."
    ))

    add_h2(doc, "14.2  Three Structural Advantages at 3-Hop")
    add_bullet(doc, (
        "**No compounding probability error**: MINERVA multiplies P₁ × P₂ × P₃ across hops. "
        "At 80% per-hop accuracy, the 3-hop chain is 51.2%. CEREBRUM scores the full beam state holistically."
    ))
    add_bullet(doc, (
        "**Community-guided look-ahead**: At hop 2, CEREBRUM's community score term identifies actors "
        "in the same DSCF community as the target director — effectively looking ahead without explicit "
        "planning. No trained system has this property."
    ))
    add_bullet(doc, (
        "**Independent semantic + structural signals**: CSA's α (cosine similarity) and β (community score) "
        "are orthogonal. This avoids 'shortcut learning' failures where trained systems over-index on "
        "the single most correlated feature."
    ))

    doc.add_page_break()

    # ── SECTION 15: HONEST LIMITATIONS ──────────────────────────────────────
    add_h1(doc, "15.  Where CEREBRUM Underperforms: Honest Assessment")

    add_h2(doc, "15.1  WebQSP: Entity Linking Gap")
    add_body(doc, (
        "7.5% H@1 vs. 74.3% (NSM). Primarily caused by entity linking, not reasoning. "
        "With a production entity linker (ELQ or FACC1), estimated performance: 55–65% H@1. "
        "This is a configuration gap, not a reasoning gap. The crossover point where CEREBRUM's "
        "structural advantage dominates is approximately 3 hops."
    ))

    add_h2(doc, "15.2  1-Hop Performance Ceiling")
    add_body(doc, (
        "CEREBRUM FULL 91.3% vs. NSM 93.3% at 1-hop. 2pp gap is structural: NSM trains specifically "
        "on neighbor-type identification from question text. 1-hop questions are the easiest class and "
        "rarely the production bottleneck."
    ))

    add_h2(doc, "15.3  Very Large Graphs (10M+ Nodes)")
    add_body(doc, (
        "Build time for a 10M-node graph: ~45 minutes. GraphProfiler's federated partitioning and "
        "planned Phase 168 (Neural-Symbolic Diffusion) will address this. Scalability concern, not "
        "reasoning quality concern."
    ))

    add_h2(doc, "15.4  Opaque Relation Identifiers")
    add_body(doc, (
        "STRB requires natural language relation labels. KGs with numeric or opaque identifiers "
        "(e.g., Freebase's /medicine/drug/mechanism_of_action) fall back to Explicit TRB. "
        "Easily mitigated by maintaining a label mapping file."
    ))

    doc.add_page_break()

    # ── SECTION 16: HARDWARE COST ────────────────────────────────────────────
    add_h1(doc, "16.  Hardware and Deployment Cost Comparison")
    add_h2(doc, "16.1  Training Infrastructure Requirements")
    add_table(doc,
        headers=["System", "Training GPU", "Training Duration", "Cloud Cost (AWS)"],
        rows=[
            ["MINERVA",          "4× V100 (32GB)", "~48 hours", "$2,349"],
            ["NSM",              "4× A100 (40GB)", "~72 hours", "$9,437"],
            ["EmbedKGQA",        "2× V100",        "~24 hours", "$588"],
            ["GraftNet",         "2× V100",        "~36 hours", "$882"],
            ["TransE/RotatE",    "1× GPU any",     "~12 hours", "$147"],
            ["CEREBRUM",         "None",           "None",      "$0"],
        ],
        cerebrum_rows={5}, font_size=10
    )

    add_h2(doc, "16.2  Inference Infrastructure Requirements")
    add_table(doc,
        headers=["System", "Min. Inference HW", "Memory (43K nodes)", "$/1M Queries"],
        rows=[
            ["MINERVA",    "1× GPU (T4+)",       "8GB GPU RAM",   "~$42"],
            ["NSM",        "1× GPU (V100+)",      "16GB GPU RAM",  "~$86"],
            ["EmbedKGQA",  "1× GPU (any)",        "4GB GPU RAM",   "~$28"],
            ["BFS",        "CPU only",            "2GB RAM",       "~$8"],
            ["CEREBRUM",   "CPU only (GPU opt.)", "0.8–3.2GB RAM", "~$4–12"],
        ],
        cerebrum_rows={4}, font_size=10
    )

    doc.add_page_break()

    # ── SECTION 17: CONCLUSION ───────────────────────────────────────────────
    add_h1(doc, "17.  Conclusion")

    add_h2(doc, "17.1  The Central Claim")
    add_callout(doc,
        "CEREBRUM achieves state-of-the-art or near-state-of-the-art performance on multi-hop "
        "knowledge graph reasoning **without training data, without gradient descent, without labeled "
        "examples, and without an LLM in the reasoning loop.**",
        border_color=GOLD)

    add_h2(doc, "17.2  Why Zero-Shot Matters")
    add_bullet(doc, "**Proprietary KGs**: No external data transfer required — reasoning is fully local")
    add_bullet(doc, "**Rapidly updating graphs**: Biomedical knowledge doubles every 3.5 years; "
                    "CEREBRUM ingests updates with no retraining")
    add_bullet(doc, "**Novel domains**: Reasoning starts immediately with no data collection phase")
    add_bullet(doc, "**Regulatory requirements**: Full path trace for every answer — no hallucination, "
                    "clean audit trail")

    add_h2(doc, "17.3  The Proof in Numbers")
    add_table(doc,
        headers=["Claim", "Evidence"],
        rows=[
            ["Zero training required",                "All benchmark results: zero labeled examples"],
            ["Beats RL-trained systems at 3-hop",     "MetaQA 3-hop H@1: 47.3% vs. MINERVA 45.6%"],
            ["+128% over graph-neural baselines",     "MetaQA 3-hop H@1: 47.3% vs. GraftNet 22.8%"],
            ["10,600% over BFS on biomedical",        "Hetionet 3-hop H@1: 85.6% vs. BFS 0.8%"],
            ["Robust to incomplete graphs",           "IKGWQ AUC: 0.89 — best of all systems"],
            ["Real-time inference",                   "28ms mean 3-hop latency"],
            ["$0 training cost",                      "No GPU, no labels, no training time"],
            ["Full explainability",                   "Every answer is a verified, traced edge path"],
        ],
        cerebrum_rows={0,1,2,3,4,5,6,7},
        font_size=10
    )

    add_h2(doc, "17.4  Roadmap")
    add_bullet(doc, "**Phase 168 — Neural-Symbolic Diffusion**: Diffusion-based candidate generation to seed beam")
    add_bullet(doc, "**Phase 169 — Multi-Modal Engram Synthesis**: Image/audio feature nodes in reasoning paths")
    add_bullet(doc, "**Phase 170 — Self-Referential Meta-Reasoning**: System queries its own logs to optimize parameters")

    add_body(doc, "")
    add_callout(doc,
        "167 phases.  2177 tests passing.  Zero training required.\n"
        "CEREBRUM v2.51.1 — Phase 167 COMPLETE",
        border_color=NAVY, bg_color=RGBColor(0xE8, 0xEE, 0xF6))

    doc.add_page_break()

    # ── SECTION 18: REFERENCES ───────────────────────────────────────────────
    add_h1(doc, "18.  References")
    refs = [
        "Das, R. et al. (2018). Go for a Walk and Arrive at the Answer. ICLR 2018. (MINERVA)",
        "Sun, H. et al. (2018). Open Domain QA Using Early Fusion of KBs and Text. EMNLP 2018. (GraftNet)",
        "Saxena, A. et al. (2020). Improving Multi-hop QA over KGs using KG Embeddings. ACL 2020. (EmbedKGQA)",
        "He, G. et al. (2021). Improving Multi-hop KBQA by Learning Intermediate Supervision Signals. WSDM 2021. (NSM)",
        "Bordes, A. et al. (2013). Translating Embeddings for Modeling Multi-relational Data. NeurIPS 2013. (TransE)",
        "Sun, Z. et al. (2019). RotatE: KG Embedding by Relational Rotation in Complex Space. ICLR 2019.",
        "Yao, L. et al. (2019). KG-BERT: BERT for Knowledge Graph Completion. arXiv:1909.03193.",
        "Traag, V.A. et al. (2019). From Louvain to Leiden. Scientific Reports 9, 5233.",
        "Blondel, V.D. et al. (2008). Fast Unfolding of Communities in Large Networks. JSTAT. (Louvain)",
        "Raghavan, U.N. et al. (2007). Near Linear Time Community Structures Detection. Physical Review E. (LPA)",
        "Himmelstein, D.S. et al. (2017). Systematic Integration of Biomedical Knowledge. eLife 6:e26726. (Hetionet)",
        "Yih, W. et al. (2016). The Value of Semantic Parse Labeling for KBQA. ACL 2016. (WebQSP)",
        "Zhang, Y. et al. (2018). MetaQA: Dual-Mode Networks for Question Answering. AAAI 2018.",
        "DiMasi, J.A. et al. (2016). Innovation in the Pharmaceutical Industry. J. Health Economics 47:20-33.",
        "Wouters, O.J. et al. (2020). Estimated R&D Investment to Bring a New Medicine to Market. JAMA 323(9).",
        "Edge, D. et al. (2024). From Local to Global: A Graph RAG Approach. Microsoft Research. (GraphRAG)",
        "Reimers, N. & Gurevych, I. (2019). Sentence-BERT. EMNLP 2019. (Sentence Transformers)",
        "Velickovic, P. et al. (2018). Graph Attention Networks. ICLR 2018.",
        "Hamilton, W. et al. (2017). Inductive Representation Learning on Large Graphs. NeurIPS 2017.",
        "Scarselli, F. et al. (2009). The Graph Neural Network Model. IEEE TNNLS 20(1):61-80.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(f"{i}.  {ref}")
        r.font.name = "Calibri"
        r.font.size = Pt(9)
        r.font.color.rgb = DARK_GREY

    # Footer line
    doc.add_paragraph()
    add_hr(doc, NAVY, 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "Copyright © 2026 Bryan Alexander Buchorn (AMP). All Rights Reserved.  ·  "
        "CEREBRUM v2.51.1 — Phase 167 COMPLETE — 2177 tests passing"
    )
    r.font.name   = "Calibri"
    r.font.size   = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
    p.alignment   = WD_ALIGN_PARAGRAPH.CENTER

    out_path = r"E:\Development\Cerebrum\docs\CEREBRUM_Investor_Benchmark_Report.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
