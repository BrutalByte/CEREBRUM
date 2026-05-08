#!/usr/bin/env python3
"""
Rebuild CEREBRUM_MASTER_SCIENTIFIC_MANUSCRIPT.docx from the Markdown source.

Converts raw Markdown/LaTeX text into a clean, readable academic Word document.
Run from the repo root:  python scripts/build_manuscript.py
"""
import re
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ──────────────────────────────────────────────────────────────────
NAVY   = (0x0D, 0x1B, 0x2A)
BLUE   = (0x1A, 0x54, 0x9A)
LTBLUE = (0x2E, 0x75, 0xB6)
SLBLUE = (0x70, 0xA8, 0xD8)
MGRAY  = (0x60, 0x60, 0x60)
LGRAY  = (0xF5, 0xF5, 0xF5)
CODEBG = (0xF0, 0xF0, 0xF0)
CODEFG = (0x30, 0x30, 0x30)
THDR   = (0x1F, 0x49, 0x7D)
BLACK  = (0x00, 0x00, 0x00)
WHITE  = (0xFF, 0xFF, 0xFF)
GOLD   = (0xC9, 0xA0, 0x2C)

def rgb(t):
    return RGBColor(t[0], t[1], t[2])

# ── LaTeX → Unicode ───────────────────────────────────────────────────────────
LATEX_SUBS = [
    # Calligraphic (before generic \mathcal)
    (r'\\mathcal\{L\}', 'ℒ'), (r'\\mathcal\{F\}', 'ℱ'),
    (r'\\mathcal\{N\}', '𝒩'), (r'\\mathcal\{G\}', '𝒢'),
    (r'\\mathcal\{S\}', '𝒮'), (r'\\mathcal\{H\}', 'ℋ'),
    (r'\\mathcal\{([A-Z])\}', r'\1'),
    # Blackboard bold
    (r'\\mathbb\{R\}', 'ℝ'), (r'\\mathbb\{N\}', 'ℕ'),
    (r'\\mathbb\{Z\}', 'ℤ'), (r'\\mathbb\{([A-Z])\}', r'\1'),
    # Greek lowercase (must be before \sigma etc. matching)
    (r'\\alpha\b',   'α'), (r'\\beta\b',    'β'), (r'\\gamma\b',   'γ'),
    (r'\\delta\b',   'δ'), (r'\\varepsilon','ε'), (r'\\epsilon\b', 'ε'),
    (r'\\zeta\b',    'ζ'), (r'\\eta\b',     'η'), (r'\\theta\b',   'θ'),
    (r'\\iota\b',    'ι'), (r'\\kappa\b',   'κ'), (r'\\lambda\b',  'λ'),
    (r'\\mu\b',      'μ'), (r'\\nu\b',      'ν'), (r'\\xi\b',      'ξ'),
    (r'\\pi\b',      'π'), (r'\\rho\b',     'ρ'), (r'\\sigma\b',   'σ'),
    (r'\\tau\b',     'τ'), (r'\\upsilon\b', 'υ'), (r'\\varphi\b',  'φ'),
    (r'\\phi\b',     'φ'), (r'\\chi\b',     'χ'), (r'\\psi\b',     'ψ'),
    (r'\\omega\b',   'ω'),
    # Greek uppercase
    (r'\\Gamma\b',   'Γ'), (r'\\Delta\b',   'Δ'), (r'\\Theta\b',   'Θ'),
    (r'\\Lambda\b',  'Λ'), (r'\\Xi\b',      'Ξ'), (r'\\Pi\b',      'Π'),
    (r'\\Sigma\b',   'Σ'), (r'\\Upsilon\b', 'Υ'), (r'\\Phi\b',     'Φ'),
    (r'\\Psi\b',     'Ψ'), (r'\\Omega\b',   'Ω'),
    # Math envs (strip wrappers)
    (r'\\begin\{aligned\}', ''), (r'\\end\{aligned\}', ''),
    (r'\\begin\{cases\}',   '{ '), (r'\\end\{cases\}', ''),
    (r'\\\\',               ' | '), (r'\s*&\s*',        '  '),
    # Decorated letters
    (r'\\widehat\{([^}]+)\}', r'(\1)̂'), (r'\\hat\{([^}]+)\}',   r'\1̂'),
    (r'\\tilde\{([^}]+)\}',   r'\1̃'),   (r'\\bar\{([^}]+)\}',   r'\1̄'),
    (r'\\vec\{([^}]+)\}',     r'\1⃗'),
    # Text / math font wrappers
    (r'\\text\{([^}]+)\}',        r'\1'),
    (r'\\mathbf\{([^}]+)\}',      r'\1'),
    (r'\\mathrm\{([^}]+)\}',      r'\1'),
    (r'\\mathit\{([^}]+)\}',      r'\1'),
    (r'\\operatorname\{([^}]+)\}',r'\1'),
    # Operators
    (r'\\cdot\b',  '·'),  (r'\\times\b', '×'),  (r'\\div\b',  '÷'),
    (r'\\leq\b',   '≤'),  (r'\\geq\b',   '≥'),  (r'\\le\b',   '≤'),
    (r'\\ge\b',    '≥'),  (r'\\neq\b',   '≠'),  (r'\\approx\b','≈'),
    (r'\\equiv\b', '≡'),  (r'\\sim\b',   '~'),
    (r'\\in\b',    ' ∈ '),(r'\\notin\b', ' ∉ '),(r'\\subset\b',' ⊂ '),
    (r'\\subseteq\b',' ⊆ '),(r'\\cup\b', ' ∪ '),(r'\\cap\b',  ' ∩ '),
    (r'\\rightarrow(?![a-zA-Z])',' → '),(r'\\leftarrow(?![a-zA-Z])',' ← '),
    (r'\\to(?![a-zA-Z])',' → '),(r'\\gets(?![a-zA-Z])',' ← '),
    (r'\\Rightarrow(?![a-zA-Z])',' ⇒ '),(r'\\Leftarrow(?![a-zA-Z])',' ⇐ '),
    (r'\\implies(?![a-zA-Z])',' ⇒ '),(r'\\iff(?![a-zA-Z])',' ⟺ '),
    (r'\\Leftrightarrow(?![a-zA-Z])',' ⟺ '),
    (r'\\forall\b','∀'),  (r'\\exists\b','∃'),
    (r'\\sum(?![a-zA-Z])','Σ'),(r'\\prod(?![a-zA-Z])','Π'),(r'\\int(?![a-zA-Z])','∫'),
    (r'\\partial\b','∂'), (r'\\nabla\b', '∇'),
    (r'\\infty\b', '∞'),  (r'\\pm\b',    '±'),  (r'\\mp\b',   '∓'),
    (r'\\ldots\b', '…'),  (r'\\cdots\b', '⋯'),  (r'\\vdots\b','⋮'),
    # Functions (use (?![a-zA-Z]) not \b — \b fails before _ and subscripts)
    (r'\\arg\\max(?![a-zA-Z])','argmax'), (r'\\arg\\min(?![a-zA-Z])','argmin'),
    (r'\\max(?![a-zA-Z])','max'), (r'\\min(?![a-zA-Z])','min'),
    (r'\\log(?![a-zA-Z])','log'), (r'\\ln(?![a-zA-Z])','ln'),
    (r'\\exp(?![a-zA-Z])','exp'), (r'\\det(?![a-zA-Z])','det'),
    (r'\\cos(?![a-zA-Z])','cos'), (r'\\sin(?![a-zA-Z])','sin'),
    (r'\\tan(?![a-zA-Z])','tan'), (r'\\tanh(?![a-zA-Z])','tanh'),
    (r'\\land\b',' ∧ '), (r'\\lor\b',' ∨ '), (r'\\lnot\b','¬'),
    (r'\\neg\b','¬'), (r'\\top\b','⊤'), (r'\\bot\b','⊥'),
    (r'\\propto\b',' ∝ '), (r'\\sim\b',' ~ '),
    (r'\\setminus\b',' \\ '), (r'\\gg\b',' ≫ '), (r'\\ll\b',' ≪ '),
    (r'\\mathbin\\[|]',' ‖ '), (r'\\mathbin\|',' ‖ '), (r'\\mathbin\\|',' ‖ '),
    (r'\\dots\b','…'), (r'\\vdash\b',' ⊢ '), (r'\\models\b',' ⊨ '),
    (r'\\perp\b','⊥'), (r'\\parallel\b','∥'),
    (r'\\mid\b',' | '), (r'\\vert\b','|'), (r'\\Vert\b','‖'),
    (r'\\left\|','|'), (r'\\right\|','|'),
    (r'\\langle\b','⟨'), (r'\\rangle\b','⟩'),
    (r'\\lceil\b','⌈'), (r'\\rceil\b','⌉'),
    (r'\\lfloor\b','⌊'), (r'\\rfloor\b','⌋'),
    (r'\\text\{sigmoid\}','σ'),
    # Citations
    (r'\\cite\{([^}]+)\}', r'[\1]'),
    # Spacing
    (r'\\quad\b','  '), (r'\\qquad\b','    '),
    (r'\\,', ' '), (r'\\;', ' '), (r'\\!', ''),
    (r'\\ ', ' '),
    (r'\\%', '%'), (r'\\#', '#'), (r'\\&', '&'),
]

# Unicode superscript map
_SUP = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶',
        '7':'⁷','8':'⁸','9':'⁹','+':'⁺','-':'⁻','n':'ⁿ','i':'ⁱ',
        'T':'ᵀ','*':'*','k':'ᵏ'}
# Unicode subscript map
_SUB = {'0':'₀','1':'₁','2':'₂','3':'₃','4':'₄','5':'₅','6':'₆',
        '7':'₇','8':'₈','9':'₉','+':'₊','-':'₋','n':'ₙ','i':'ᵢ',
        'k':'ₖ','v':'ᵥ','u':'ᵤ','t':'ₜ'}

def _to_sup(s):
    return ''.join(_SUP.get(c, c) for c in s)

def _to_sub(s):
    if len(s) <= 2 and all(c in _SUB for c in s):
        return ''.join(_SUB[c] for c in s)
    return '_' + s

def _extract_group(text, pos):
    """Extract content of a {}-delimited group at pos, return (content, end_pos)."""
    if pos >= len(text) or text[pos] != '{':
        return '', pos
    depth = 1
    i = pos + 1
    while i < len(text) and depth:
        if text[i] == '{':
            depth += 1
        elif text[i] == '}':
            depth -= 1
        i += 1
    return text[pos+1:i-1], i

def _expand_fracs(text):
    """Replace \\frac{num}{den} → (num)/(den), handling nested braces."""
    out = []
    i = 0
    while i < len(text):
        if text[i:i+5] == r'\frac' and i+5 < len(text) and text[i+5] == '{':
            num, after_num = _extract_group(text, i+5)
            if after_num < len(text) and text[after_num] == '{':
                den, after_den = _extract_group(text, after_num)
                out.append(f'({num})/({den})')
                i = after_den
                continue
        out.append(text[i])
        i += 1
    return ''.join(out)

def _expand_wrappers(text):
    """Expand LaTeX wrapping commands that contain {} groups."""
    # Calligraphic
    text = re.sub(r'\\mathcal\{([A-Z])\}', lambda m: {
        'L':'ℒ','F':'ℱ','N':'𝒩','G':'𝒢','S':'𝒮','H':'ℋ','C':'𝒞','R':'ℛ'
    }.get(m.group(1), m.group(1)), text)
    # Blackboard
    text = re.sub(r'\\mathbb\{([A-Z])\}', lambda m: {
        'R':'ℝ','N':'ℕ','Z':'ℤ','Q':'ℚ','C':'ℂ'
    }.get(m.group(1), m.group(1)), text)
    # Text/math font wrappers — preserve content, discard wrapper
    for cmd in (r'\\text', r'\\mathbf', r'\\mathrm', r'\\mathit',
                r'\\textbf', r'\\textit', r'\\texttt',
                r'\\operatorname', r'\\boldsymbol'):
        text = re.sub(cmd + r'\{([^{}]+)\}', r'\1', text)
    # Decorated letters
    text = re.sub(r'\\widehat\{([^}]+)\}', r'(\1)̂', text)
    text = re.sub(r'\\hat\{([^}]+)\}',     r'\1̂',   text)
    text = re.sub(r'\\tilde\{([^}]+)\}',   r'\1̃',   text)
    text = re.sub(r'\\bar\{([^}]+)\}',     r'\1̄',   text)
    text = re.sub(r'\\vec\{([^}]+)\}',     r'\1⃗',   text)
    text = re.sub(r'\\overline\{([^}]+)\}',r'\1̄',   text)
    # \sqrt{x}
    text = re.sub(r'\\sqrt\{([^}]+)\}',    r'√(\1)', text)
    # \left and \right (size wrappers — just remove the command)
    text = re.sub(r'\\left([({[\|])', r'\1', text)
    text = re.sub(r'\\right([)}\]|])', r'\1', text)
    text = re.sub(r'\\left\.', '', text)
    text = re.sub(r'\\right\.', '', text)
    return text

def clean_math(text):
    """Convert LaTeX math to readable Unicode."""
    # Protect \{ and \} (escaped braces — not group delimiters)
    text = text.replace(r'\{', '\u27E8').replace(r'\}', '\u27E9')
    # Remove environments
    text = re.sub(r'\\begin\{[^}]+\}', '', text)
    text = re.sub(r'\\end\{[^}]+\}', '', text)
    # Multiple passes: unwrap content commands, then expand fracs, then symbols
    for _ in range(5):
        text = _expand_wrappers(text)
        text = _expand_fracs(text)
        # Apply substitution table each pass so operators resolve before subscripts
        for pat, repl in LATEX_SUBS:
            text = re.sub(pat, repl, text)
        # ^{...} superscripts
        text = re.sub(r'\^\{([^}]+)\}', lambda m: _to_sup(m.group(1)), text)
        text = re.sub(r'\^([a-zA-Z0-9*])', lambda m: _to_sup(m.group(1)), text)
        # _{...} subscripts
        text = re.sub(r'_\{([^}]+)\}', lambda m: _to_sub(m.group(1)), text)
        text = re.sub(r'_([a-zA-Z0-9])', lambda m: _to_sub(m.group(1)), text)
    # Final pass — strip remaining bare braces and restore escaped parens
    text = re.sub(r'[{}]', '', text)
    text = text.replace('\u27E8', '(').replace('\u27E9', ')')
    text = re.sub(r' {3,}', '  ', text)
    return text.strip()

# ── Inline tokenizer ─────────────────────────────────────────────────────────
def tokenize(text):
    """Yield (kind, content) tokens from inline Markdown/LaTeX text."""
    i = 0
    buf = []
    while i < len(text):
        # Bold: **...**
        if text[i:i+2] == '**':
            if buf:
                yield ('text', ''.join(buf)); buf = []
            end = text.find('**', i+2)
            if end != -1:
                yield ('bold', text[i+2:end]); i = end + 2; continue
        # Italic: *...* (not part of **)
        if text[i] == '*' and text[i:i+2] != '**':
            if buf:
                yield ('text', ''.join(buf)); buf = []
            end = text.find('*', i+1)
            if end != -1 and text[end:end+2] != '**':
                yield ('italic', text[i+1:end]); i = end + 1; continue
        # Inline code: `...`
        if text[i] == '`':
            if buf:
                yield ('text', ''.join(buf)); buf = []
            end = text.find('`', i+1)
            if end != -1:
                yield ('code', text[i+1:end]); i = end + 1; continue
        # Display math: $$...$$
        if text[i:i+2] == '$$':
            if buf:
                yield ('text', ''.join(buf)); buf = []
            end = text.find('$$', i+2)
            if end != -1:
                yield ('math', text[i+2:end]); i = end + 2; continue
        # Inline math: $...$
        if text[i] == '$':
            if buf:
                yield ('text', ''.join(buf)); buf = []
            end = text.find('$', i+1)
            if end != -1:
                yield ('math', text[i+1:end]); i = end + 1; continue
        buf.append(text[i]); i += 1
    if buf:
        yield ('text', ''.join(buf))

# ── Document helpers ─────────────────────────────────────────────────────────
def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(*color)
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def para_shade(para, color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(*color)
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)

def add_hrule(doc):
    """Add a thin horizontal rule as a paragraph bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '{:02X}{:02X}{:02X}'.format(*SLBLUE))
    pBdr.append(bot); pPr.append(pBdr)

def page_break_before(para):
    pPr = para._p.get_or_add_pPr()
    pb  = OxmlElement('w:pageBreakBefore')
    pb.set(qn('w:val'), '1')
    pPr.append(pb)

def add_run(para, text, bold=False, italic=False, size=None,
            color=None, font='Cambria', mono=False):
    if not text:
        return
    run = para.add_run(text)
    run.font.name = 'Courier New' if mono else font
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    run.font.bold   = bold
    run.font.italic = italic

def preprocess(text):
    """Clean LaTeX commands that appear in plain text (outside $ markers).
    Applies full clean_math() conversion to any backslash-command sequences,
    and also converts _{...} / ^{...} subscript/superscript notation."""
    # Always convert _{...} and ^{...} subscript/superscript even without backslash
    text = re.sub(r'_\{([^}]+)\}', lambda m: _to_sub(m.group(1)), text)
    text = re.sub(r'\^\{([^}]+)\}', lambda m: _to_sup(m.group(1)), text)
    if '\\' not in text:
        return text
    # Run the same multi-pass conversion used for display math
    text = text.replace(r'\{', '\u27E8').replace(r'\}', '\u27E9')
    text = re.sub(r'\\begin\{[^}]+\}', '', text)
    text = re.sub(r'\\end\{[^}]+\}', '', text)
    for _ in range(3):
        text = _expand_wrappers(text)
        text = _expand_fracs(text)
        for pat, repl in LATEX_SUBS:
            text = re.sub(pat, repl, text)
        text = re.sub(r'\^\{([^}]+)\}', lambda m: _to_sup(m.group(1)), text)
        text = re.sub(r'\^([a-zA-Z0-9*])', lambda m: _to_sup(m.group(1)), text)
        text = re.sub(r'_\{([^}]+)\}', lambda m: _to_sub(m.group(1)), text)
        text = re.sub(r'_([a-zA-Z0-9])', lambda m: _to_sub(m.group(1)), text)
    text = re.sub(r'[{}]', '', text)
    text = text.replace('\u27E8', '(').replace('\u27E9', ')')
    return text

def add_inline(para, text, base_size=11, base_font='Cambria', base_color=BLACK, base_bold=False):
    """Parse inline Markdown/math tokens and add runs to paragraph."""
    text = preprocess(text)
    for kind, content in tokenize(text):
        if kind == 'text':
            add_run(para, content, bold=base_bold, size=base_size, font=base_font, color=base_color)
        elif kind == 'bold':
            add_run(para, content, bold=True, size=base_size, font=base_font, color=base_color)
        elif kind == 'italic':
            add_run(para, content, italic=True, size=base_size, font=base_font, color=base_color)
        elif kind == 'code':
            r = para.add_run(content)
            r.font.name = 'Courier New'
            r.font.size = Pt(base_size - 1)
            r.font.color.rgb = rgb(CODEFG)
        elif kind == 'math':
            cleaned = clean_math(content)
            add_run(para, cleaned, bold=base_bold, italic=True, size=base_size,
                    font=base_font, color=BLUE)

def heading(doc, text, level):
    """Add a styled heading paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt({1:18, 2:14, 3:10, 4:8}.get(level, 8))
    p.paragraph_format.space_after  = Pt({1:10, 2:6,  3:4,  4:3}.get(level, 3))
    p.paragraph_format.keep_with_next = True
    if level == 1:
        page_break_before(p)
        add_inline(p, text, base_size=18, base_color=NAVY, base_bold=True)
    elif level == 2:
        add_inline(p, text, base_size=13, base_color=NAVY, base_bold=True)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '2')
        bot.set(qn('w:color'), '{:02X}{:02X}{:02X}'.format(*BLUE))
        pBdr.append(bot); pPr.append(pBdr)
    elif level == 3:
        add_inline(p, text, base_size=12, base_color=BLUE, base_bold=True)
    else:
        add_inline(p, text, base_size=11, base_color=LTBLUE, base_bold=True)
    return p

def body(doc, text, indent=0):
    """Add a body paragraph with inline markup."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    add_inline(p, text)
    return p

def math_block(doc, text):
    """Add a display math paragraph."""
    cleaned = clean_math(text)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_shade(p, (0xEE, 0xF4, 0xFB))
    r = p.add_run(cleaned)
    r.font.name  = 'Cambria'
    r.font.size  = Pt(11)
    r.font.color.rgb = rgb(NAVY)
    r.font.italic = True
    return p

def code_block(doc, lines, lang=''):
    """Add a code block with monospace font and gray background."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.2)
        para_shade(p, CODEBG)
        r = p.add_run(line if line else ' ')
        r.font.name  = 'Courier New'
        r.font.size  = Pt(9)
        r.font.color.rgb = rgb(CODEFG)
    # spacing after block
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after  = Pt(6)

def bullet(doc, text, level=0, ordered=False, num=None):
    """Add a bullet/numbered list item."""
    p = doc.add_paragraph()
    indent_base = 0.3
    p.paragraph_format.left_indent  = Inches(indent_base + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    # Bullet marker
    if ordered and num:
        marker = f'{num}. '
    else:
        markers = ['•', '◦', '▪']
        marker = markers[min(level, 2)] + '  '
    add_run(p, marker, bold=(level == 0), size=11, color=BLUE if level == 0 else MGRAY)
    add_inline(p, text)
    return p

def metadata_line(doc, text):
    """Add an author/status/date line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    add_inline(p, text, base_size=10, base_color=MGRAY)
    return p

def ref_line(doc, text):
    """Add a reference list entry."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_before      = Pt(2)
    p.paragraph_format.space_after       = Pt(2)
    # Strip leading number/bullet
    text = re.sub(r'^\s*[-*]\s*', '', text)
    add_inline(p, text, base_size=9.5, base_color=MGRAY)
    return p

def add_md_table(doc, rows):
    """Render a parsed Markdown table as a Word table."""
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci >= col_count:
                break
            cell = tbl.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            if ri == 0:
                set_cell_bg(cell, THDR)
                add_run(p, cell_text.strip(), bold=True, size=9.5,
                        color=WHITE, font='Cambria')
            else:
                bg = LGRAY if ri % 2 == 0 else WHITE
                set_cell_bg(cell, bg)
                add_inline(p, cell_text.strip())
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Markdown table parser ─────────────────────────────────────────────────────
def parse_table_row(line):
    """Parse a markdown table row into a list of cell strings."""
    cells = [c.strip() for c in line.strip().strip('|').split('|')]
    return cells

def is_table_sep(line):
    """True if this is a markdown table separator row like |---|---|"""
    stripped = line.strip()
    return bool(re.match(r'^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)*\|?$', stripped))

# ── Page setup ───────────────────────────────────────────────────────────────
def setup_doc():
    doc = Document()
    for section in doc.sections:
        section.page_width   = Inches(8.5)
        section.page_height  = Inches(11)
        section.left_margin  = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin   = Inches(1.0)
        section.bottom_margin= Inches(1.0)
    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name  = 'Cambria'
    style.font.size  = Pt(11)
    style.paragraph_format.space_after  = Pt(4)
    style.paragraph_format.line_spacing = Pt(14)
    return doc

# ── Main parser ───────────────────────────────────────────────────────────────
def build(src_path, out_path):
    doc = setup_doc()

    # Cover page
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(48)
    add_run(cover, 'CEREBRUM', bold=True, size=28, color=NAVY)
    cover.add_run('\n')
    add_run(cover, 'Master Scientific Manuscript', bold=False, size=16, color=BLUE)
    cover.add_run('\n')
    add_run(cover, 'v2.51.1 · Phase 167 · 37 Research Papers', size=11, color=MGRAY)
    doc.add_paragraph().add_run().add_break(
        __import__('docx.enum.text', fromlist=['WD_BREAK_TYPE']).WD_BREAK_TYPE.PAGE
    )

    with open(src_path, encoding='utf-8') as f:
        lines = f.readlines()

    # State
    in_code   = False
    in_math   = False
    code_buf  = []
    math_buf  = []
    code_lang = ''
    tbl_buf   = []           # pending table rows
    in_table  = False
    is_ref_section  = False
    is_meta_section = False  # right after paper title (authors/status/date)
    list_counters = {}       # level → count for ordered lists

    def flush_table():
        nonlocal tbl_buf, in_table
        if tbl_buf:
            add_md_table(doc, tbl_buf)
        tbl_buf = []; in_table = False

    def flush_code():
        nonlocal code_buf, in_code
        code_block(doc, code_buf, code_lang)
        code_buf = []; in_code = False

    for raw in lines:
        line = raw.rstrip('\n')
        stripped = line.strip()

        # ── Code block ─────────────────────────────────────────────────────
        if stripped.startswith('```'):
            if in_table: flush_table()
            if in_code:
                flush_code()
            else:
                in_code = True
                code_lang = stripped[3:].strip()
                code_buf = []
            continue

        if in_code:
            code_buf.append(line)
            continue

        # ── Display math block ─────────────────────────────────────────────
        if stripped.startswith('$$') or (in_math and stripped.endswith('$$')):
            if in_table: flush_table()
            if in_math:
                # Closing marker: either bare $$ or content$$
                tail = stripped[:-2].strip() if stripped.endswith('$$') and stripped != '$$' else ''
                if tail:
                    math_buf.append(tail)
                math_block(doc, '\n'.join(math_buf))
                math_buf = []; in_math = False
            else:
                # Single-line $$...$$ on one line
                if stripped.startswith('$$') and stripped.endswith('$$') and len(stripped) > 4:
                    math_block(doc, stripped[2:-2])
                else:
                    # Start of multi-line block: $$  or  $$\begin{aligned}
                    in_math = True
                    head = stripped[2:].strip()
                    math_buf = [head] if head else []
            continue

        if in_math:
            math_buf.append(stripped)
            continue

        # ── Table detection ────────────────────────────────────────────────
        if stripped.startswith('|') and stripped.endswith('|'):
            if is_table_sep(stripped):
                continue  # skip separator row
            tbl_buf.append(parse_table_row(stripped))
            in_table = True
            continue
        elif in_table:
            flush_table()

        # ── Empty line ─────────────────────────────────────────────────────
        if not stripped:
            is_meta_section = False
            continue

        # ── Headings ───────────────────────────────────────────────────────
        if stripped.startswith('#'):
            m = re.match(r'^(#{1,4})\s+(.*)', stripped)
            if m:
                level = len(m.group(1))
                title = m.group(2).strip()
                if level == 1:
                    is_ref_section  = False
                    is_meta_section = True
                    list_counters   = {}
                heading(doc, title, level)
                # Detect references section
                if re.match(r'^(references|bibliography)', title, re.I):
                    is_ref_section = True
                continue

        # ── Horizontal rule ────────────────────────────────────────────────
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            add_hrule(doc)
            is_meta_section = False
            continue

        # ── Ordered list ───────────────────────────────────────────────────
        om = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if om:
            indent_ws = om.group(1)
            num       = int(om.group(2))
            text      = om.group(3)
            level     = len(indent_ws) // 2
            bullet(doc, text, level=level, ordered=True, num=num)
            continue

        # ── Unordered list ─────────────────────────────────────────────────
        ulm = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if ulm:
            indent_ws = ulm.group(1)
            text      = ulm.group(2)
            level     = len(indent_ws) // 2
            bullet(doc, text, level=level)
            continue

        # ── Metadata lines (right after paper title) ───────────────────────
        if is_meta_section and re.match(r'^\*\*\w', stripped):
            metadata_line(doc, stripped)
            continue

        # ── References section ─────────────────────────────────────────────
        if is_ref_section:
            if re.match(r'^(\d+\.|[-*])\s', stripped):
                ref_line(doc, stripped)
                continue

        # ── Copyright / Reviewed lines ─────────────────────────────────────
        if re.match(r'^\*\*(Copyright|Reviewed)', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            add_inline(p, stripped, base_size=9, base_color=MGRAY)
            continue

        # ── Single-line display math: $$...$$  ─────────────────────────────
        if stripped.startswith('$$') and stripped.endswith('$$'):
            math_block(doc, stripped[2:-2])
            continue

        # ── Default: body paragraph ────────────────────────────────────────
        body(doc, stripped)

    # Flush any open blocks
    if in_code:   flush_code()
    if in_math:   math_block(doc, '\n'.join(math_buf))
    if in_table:  flush_table()

    doc.save(out_path)
    para_count  = len(doc.paragraphs)
    table_count = len(doc.tables)
    print(f'OK Saved {out_path}')
    print(f'   {para_count} paragraphs, {table_count} tables')


if __name__ == '__main__':
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    src  = os.path.join(root, 'docs', 'CEREBRUM_MASTER_SCIENTIFIC_MANUSCRIPT.md')
    out  = os.path.join(root, 'docs', 'CEREBRUM_MASTER_SCIENTIFIC_MANUSCRIPT.docx')
    build(src, out)
