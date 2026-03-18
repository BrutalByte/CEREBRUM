"""
Generate the full Parallax genesis conversation transcript as a Word document.
Covers from the first cluster animation request through the white paper completion.
"""
import json, re
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

JSONL_PATH  = r"C:\Users\bryan\.claude\projects\E--Development-AURA\2cdef21d-f550-41db-a96f-175b1c7bb1cc.jsonl"
OUTPUT_PATH = r"C:\Users\bryan\Documents\Parallax_Genesis_Transcript.docx"

TITLE    = "Parallax: The Genesis Conversation"
SUBTITLE = "Complete Session Transcript with Annotations"
AUTHOR   = "Bryan Alexander Buchorn  \u00b7  AMP"
AFFIL    = "Independent Researcher"
EMAIL    = "bryan.alexander@buchorn.com"
DATE     = "March 2026"

COLOR_USER   = RGBColor(0,   70,  127)   # deep blue  — user messages
COLOR_ASST   = RGBColor(30,  30,  30)    # near-black — assistant messages
COLOR_TOOL   = RGBColor(100, 100, 100)   # gray       — tool calls
COLOR_TS     = RGBColor(130, 130, 130)   # light gray — timestamps
COLOR_ANNOT  = RGBColor(160, 80,  0)     # amber      — editorial annotations
COLOR_USER_BG = "E8F4FD"
COLOR_ASST_BG = "F5F5F5"
COLOR_TOOL_BG = "F0F0F0"


# ── Annotation map: message index → annotation text ─────────────────────────
ANNOTATIONS = {
    138: "GENESIS MOMENT: The session that would become Parallax begins here. "
         "A live engineering request — animate the clustering process in real-time — "
         "is about to open a chain of questions that leads to a novel theoretical "
         "framework for knowledge graph reasoning.",

    194: "ALGORITHM SWITCH: The Louvain-to-Leiden migration. This is the first "
         "moment the conversation focuses deeply on what community detection "
         "algorithms actually compute and why their differences matter.",

    229: "KEY COMPARISON: The first structured analysis of Louvain vs Leiden vs LPA. "
         "Understanding these three algorithms side-by-side is the prerequisite "
         "for the question that follows.",

    231: "SIMULTANEITY QUESTION: 'Can the algorithm include structure from both?' "
         "This is the inflection point. The question is no longer about which "
         "algorithm to use sequentially — it is asking whether a fundamentally "
         "different kind of algorithm is possible. DSCF is born here.",
}

# Find the user messages that mark pivotal moments by content substring
PIVOT_MARKERS = [
    ("real-time", "GENESIS: First request for real-time cluster animation"),
    ("Louvian Algorhith", "PIVOT: Request to switch from Louvain to Leiden"),
    ("add the ability to switch", "EXPANSION: Adding LPA as alternative algorithm"),
    ("stacked or run one after another", "CRITICAL QUESTION: Sequential vs. simultaneous algorithms"),
    ("both simultaneously", "DSCF BIRTH: The question that creates a novel algorithm"),
    ("new methodology", "THEORETICAL LEAP: 'Is that a new methodology?'"),
    ("treat knowledge", "THE BIG QUESTION: 'How can we treat Knowledge Graphs like LLMs?'"),
    ("let's try to figure it out", "COMMITMENT: 'Let's make history!'"),
    ("not implement it just yet", "STRATEGIC DECISION: Spin off as separate framework-agnostic project"),
    ("white paper", "DOCUMENTATION: Request for white paper format"),
]


def add_border_bottom(p, color="AAAAAA", sz="4"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), sz)
    b.set(qn("w:space"), "1"); b.set(qn("w:color"), color)
    pBdr.append(b); pPr.append(pBdr)


def add_border_left(p, color="4466A0", sz="18"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "6"); left.set(qn("w:color"), color)
    pBdr.append(left); pPr.append(pBdr)


def set_para_shading(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def setup_styles(doc):
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)
    for level, size, sb in [(1, 14, 16), (2, 11, 10), (3, 10, 8)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"; h.font.size = Pt(size); h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(sb)
        h.paragraph_format.space_after  = Pt(4)
        h.paragraph_format.keep_with_next = True


def fmt_ts(ts_str):
    """Format ISO timestamp to human-readable."""
    try:
        dt = datetime.fromisoformat(ts_str.replace("Z", "+00:00"))
        return dt.strftime("%Y-%m-%d  %H:%M:%S UTC")
    except:
        return ts_str


def fmt_ts_short(ts_str):
    try:
        dt = datetime.fromisoformat(ts_str.replace("Z", "+00:00"))
        return dt.strftime("%H:%M:%S")
    except:
        return ts_str


def extract_messages(path):
    entries = []
    with open(path, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line: continue
            try: entries.append(json.loads(line))
            except: pass

    msgs = []
    for e in entries:
        if e.get("type") not in ("user", "assistant"): continue
        ts   = e.get("timestamp", "")
        role = e.get("type")
        content = e.get("message", {}).get("content", "")
        model = e.get("message", {}).get("model", "")

        text_parts = []
        tool_parts = []

        if isinstance(content, str):
            text_parts.append(content)
        elif isinstance(content, list):
            for block in content:
                if not isinstance(block, dict): continue
                btype = block.get("type", "")
                if btype == "text":
                    t = block.get("text", "").strip()
                    if t: text_parts.append(t)
                elif btype == "thinking":
                    pass  # omit
                elif btype == "tool_use":
                    name = block.get("name", "tool")
                    inp  = block.get("input", {})
                    if name == "Bash":
                        desc = inp.get("description", "")
                        cmd  = inp.get("command", "")[:300].replace("\n", " | ")
                        display = desc if desc else cmd
                        tool_parts.append(f"Bash: {display}")
                    elif name == "Read":
                        tool_parts.append(f"Read: {inp.get('file_path','')}")
                    elif name == "Write":
                        tool_parts.append(f"Write: {inp.get('file_path','')}")
                    elif name == "Edit":
                        fp = inp.get("file_path", "")
                        old = (inp.get("old_string","")[:60] + "...").replace("\n"," ")
                        tool_parts.append(f"Edit: {fp}  [{old}]")
                    elif name == "Glob":
                        tool_parts.append(f"Glob: {inp.get('pattern','')}")
                    elif name == "Grep":
                        tool_parts.append(f"Grep: {inp.get('pattern','')} in {inp.get('path','.')}")
                    elif name == "Agent":
                        tool_parts.append(f"Agent ({inp.get('subagent_type','general')}): {inp.get('description','')}")
                    elif name == "ExitPlanMode":
                        tool_parts.append("ExitPlanMode: submitted plan for approval")
                    elif name == "EnterPlanMode":
                        tool_parts.append("EnterPlanMode: entered planning mode")
                    elif name in ("TaskCreate","TaskUpdate","TaskGet","TaskList"):
                        tool_parts.append(f"{name}")
                    elif name == "AskUserQuestion":
                        qs = inp.get("questions", [])
                        q_texts = [q.get("question","") for q in qs]
                        tool_parts.append(f"AskUserQuestion: {' | '.join(q_texts)}")
                    else:
                        tool_parts.append(f"{name}")
                elif btype == "tool_result":
                    pass  # skip

        full_text = "\n\n".join(text_parts).strip()
        # Filter out system context injections (very long user messages with "session is being continued")
        if role == "user" and "session is being continued from a previous conversation" in full_text:
            full_text = "[Context window reset — conversation summary injected by system]"
            tool_parts = []

        if full_text or tool_parts:
            msgs.append({
                "ts": ts,
                "role": role,
                "model": model,
                "text": full_text,
                "tools": tool_parts,
                "idx": len(msgs),
            })

    return msgs


def find_start(msgs):
    for i, m in enumerate(msgs):
        if m["role"] == "user" and "real-time" in m["text"].lower() and "cluster" in m["text"].lower():
            return i
    return 0


def detect_pivot(text):
    for marker, label in PIVOT_MARKERS:
        if marker.lower() in text.lower():
            return label
    return None


def write_annotation(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    add_border_left(p, color="A0701C", sz="12")
    r = p.add_run(f"\u25b6 EDITORIAL NOTE: {text}")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(140, 70, 0); r.italic = True


def write_message(doc, msg, seq_num):
    role = msg["role"]
    ts   = msg["ts"]
    text = msg["text"]
    tools = msg["tools"]

    is_user = (role == "user")
    color   = COLOR_USER if is_user else COLOR_ASST
    label   = "USER" if is_user else "ASSISTANT"
    bg      = COLOR_USER_BG if is_user else COLOR_ASST_BG

    # ── Header line ───────────────────────────────────────────────────────────
    hp = doc.add_paragraph()
    hp.paragraph_format.space_before = Pt(8)
    hp.paragraph_format.space_after  = Pt(1)
    add_border_bottom(hp, color="CCCCCC", sz="2")

    r_num = hp.add_run(f"#{seq_num:03d}  ")
    r_num.font.name = "Courier New"; r_num.font.size = Pt(8)
    r_num.font.color.rgb = COLOR_TS

    r_lbl = hp.add_run(label)
    r_lbl.font.name = "Calibri"; r_lbl.font.size = Pt(9.5)
    r_lbl.bold = True; r_lbl.font.color.rgb = color

    r_ts = hp.add_run(f"  \u2014  {fmt_ts(ts)}")
    r_ts.font.name = "Calibri"; r_ts.font.size = Pt(8.5)
    r_ts.font.color.rgb = COLOR_TS

    if msg.get("model"):
        r_m = hp.add_run(f"  [{msg['model']}]")
        r_m.font.name = "Calibri"; r_m.font.size = Pt(8)
        r_m.font.color.rgb = COLOR_TS

    # ── Text content ──────────────────────────────────────────────────────────
    if text:
        # Split into paragraphs
        paras = [p.strip() for p in text.split("\n\n") if p.strip()]
        if not paras:
            paras = [p.strip() for p in text.split("\n") if p.strip()]

        for para_text in paras[:40]:  # cap at 40 paragraphs per message
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.25)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            if is_user:
                add_border_left(p, color="1E6E9E", sz="10")

            # Light inline formatting
            pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
            for part in pattern.split(para_text):
                if part.startswith("**") and part.endswith("**") and len(part) > 4:
                    r = p.add_run(part[2:-2]); r.bold = True
                    r.font.name = "Calibri"; r.font.size = Pt(10)
                    r.font.color.rgb = color
                elif part.startswith("`") and part.endswith("`") and len(part) > 2:
                    r = p.add_run(part[1:-1])
                    r.font.name = "Courier New"; r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(149, 55, 53)
                elif part:
                    r = p.add_run(part)
                    r.font.name = "Calibri"; r.font.size = Pt(10)
                    r.font.color.rgb = color

    # ── Tool calls ────────────────────────────────────────────────────────────
    if tools:
        for tool_line in tools:
            tp = doc.add_paragraph()
            tp.paragraph_format.left_indent  = Inches(0.35)
            tp.paragraph_format.space_before = Pt(1)
            tp.paragraph_format.space_after  = Pt(1)
            r = tp.add_run(f"  \u2192 {tool_line}")
            r.font.name = "Courier New"; r.font.size = Pt(8.5)
            r.font.color.rgb = COLOR_TOOL


def make_doc():
    doc = Document()
    setup_styles(doc)
    for sec in doc.sections:
        sec.page_width    = Inches(8.5); sec.page_height = Inches(11)
        sec.top_margin    = sec.bottom_margin = Inches(1.0)
        sec.left_margin   = sec.right_margin  = Inches(1.1)

    # ── Title block ───────────────────────────────────────────────────────────
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, tl in enumerate(TITLE.split("\n")):
        run = tp.add_run(tl); run.bold = True
        run.font.name = "Calibri"; run.font.size = Pt(20)
        if j < len(TITLE.split("\n")) - 1: tp.add_run("\n")
    tp.paragraph_format.space_before = Pt(0); tp.paragraph_format.space_after = Pt(6)

    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sp.add_run(SUBTITLE); rs.italic = True
    rs.font.name = "Calibri"; rs.font.size = Pt(12)
    sp.paragraph_format.space_before = Pt(0); sp.paragraph_format.space_after = Pt(10)

    ap = doc.add_paragraph(); ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = ap.add_run(AUTHOR); r1.bold = True
    r1.font.name = "Calibri"; r1.font.size = Pt(11)
    ap.paragraph_format.space_after = Pt(2)

    af = doc.add_paragraph(); af.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = af.add_run(AFFIL); r2.italic = True
    r2.font.name = "Calibri"; r2.font.size = Pt(10)
    af.paragraph_format.space_after = Pt(2)

    em = doc.add_paragraph(); em.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = em.add_run(EMAIL); r3.font.name = "Courier New"; r3.font.size = Pt(9)
    em.paragraph_format.space_after = Pt(2)

    dv = doc.add_paragraph(); dv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = dv.add_run(f"{DATE}  \u00b7  Session ID: 2cdef21d-f550-41db-a96f-175b1c7bb1cc")
    r4.italic = True; r4.font.name = "Calibri"; r4.font.size = Pt(9)
    dv.paragraph_format.space_after = Pt(12)
    add_border_bottom(dv)

    # ── Prefatory note ────────────────────────────────────────────────────────
    pn = doc.add_paragraph()
    pn.paragraph_format.space_before = Pt(10)
    pn.paragraph_format.space_after  = Pt(10)
    pn.paragraph_format.left_indent  = Inches(0.4)
    pn.paragraph_format.right_indent = Inches(0.4)
    add_border_left(pn, color="888888", sz="6")
    r = pn.add_run(
        "This document is the verbatim transcript of the Claude Code session in which the Parallax "
        "framework was conceived. It begins at the moment Bryan Alexander Buchorn (AMP) requested "
        "real-time cluster animation for the AURA knowledge graph, and continues through the completion "
        "of the Parallax white paper. Editorial annotations (amber, marked with \u25b6) mark the "
        "pivotal conceptual transitions. Tool calls are shown inline in monospace. The session ran "
        "across multiple context windows; system-injected summaries are noted where they occur. "
        "All timestamps are UTC."
    )
    r.font.name = "Calibri"; r.font.size = Pt(9.5); r.italic = True
    r.font.color.rgb = RGBColor(60, 60, 60)

    add_border_bottom(pn, color="AAAAAA")

    # ── Load and filter messages ───────────────────────────────────────────────
    print("Loading transcript...")
    msgs = extract_messages(JSONL_PATH)
    start = find_start(msgs)
    msgs  = msgs[start:]
    print(f"  {len(msgs)} messages from genesis point onward")

    # ── Chapter markers ───────────────────────────────────────────────────────
    CHAPTERS = {
        0:  "Chapter 1: The Visualization Problem",
        None: None,  # dynamic based on content
    }

    chapter_markers = [
        ("real-time", "real-time",     "Chapter 1: The Visualization Problem — Real-Time Cluster Animation"),
        ("Louvian",   "Leiden",         "Chapter 2: Algorithm Exploration — From Louvain to Leiden"),
        ("switch between", "LPA",       "Chapter 3: Algorithm Choice — Adding Label Propagation"),
        ("stacked",   "simultaneously", "Chapter 4: The Simultaneity Question — Birth of DSCF"),
        ("new methodology", "LLMs",     "Chapter 5: The Theoretical Leap — Knowledge Graphs as Language Models"),
        ("let's try",  "history",       "Chapter 6: Commitment — Building the Framework"),
        ("not implement", "repository", "Chapter 7: Strategic Pivot — The Standalone Framework"),
        ("white paper", "white paper",  "Chapter 8: The White Paper — Parallax Formalized"),
    ]

    chapter_triggers = []
    for kw1, kw2, title in chapter_markers:
        for i, m in enumerate(msgs):
            t = m["text"].lower()
            if kw1.lower() in t and kw2.lower() in t:
                chapter_triggers.append((i, title))
                break

    chapter_set = {idx: title for idx, title in chapter_triggers}

    # ── Render messages ───────────────────────────────────────────────────────
    seq = 1
    prev_ts = None

    for i, msg in enumerate(msgs):

        # Chapter heading
        if i in chapter_set:
            ch = doc.add_heading(chapter_set[i], level=1)
            ch.paragraph_format.space_before = Pt(18)
            add_border_bottom(ch, color="222222", sz="6")

        # Time gap callout (> 30 minutes between messages)
        if prev_ts and msg["ts"]:
            try:
                t1 = datetime.fromisoformat(prev_ts.replace("Z","+00:00"))
                t2 = datetime.fromisoformat(msg["ts"].replace("Z","+00:00"))
                gap_min = (t2 - t1).total_seconds() / 60
                if gap_min > 30:
                    gp = doc.add_paragraph()
                    gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    rg = gp.add_run(f"\u23f1  {int(gap_min)} minutes elapsed  \u23f1")
                    rg.font.name = "Calibri"; rg.font.size = Pt(8.5)
                    rg.italic = True; rg.font.color.rgb = COLOR_TS
                    gp.paragraph_format.space_before = Pt(6)
                    gp.paragraph_format.space_after  = Pt(6)
            except:
                pass

        # Editorial annotation for pivotal user messages
        if msg["role"] == "user" and msg["text"] and msg["text"] != "[Context window reset — conversation summary injected by system]":
            pivot = detect_pivot(msg["text"])
            if pivot:
                write_annotation(doc, pivot)

        write_message(doc, msg, seq)
        seq += 1
        if msg["ts"]:
            prev_ts = msg["ts"]

    # ── Closing ───────────────────────────────────────────────────────────────
    doc.add_page_break()
    ep = doc.add_paragraph()
    ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_border_bottom(ep, color="888888")

    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(16)
    r = fp.add_run("End of Transcript")
    r.font.name = "Calibri"; r.font.size = Pt(14); r.bold = True

    sp2 = doc.add_paragraph()
    sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sp2.add_run(
        f"Session 2cdef21d-f550-41db-a96f-175b1c7bb1cc\n"
        f"{len(msgs)} messages rendered\n"
        f"AURA Project  \u00b7  E:\\Development\\AURA\n"
        f"Bryan Alexander Buchorn (AMP)  \u00b7  {DATE}"
    )
    r2.font.name = "Calibri"; r2.font.size = Pt(9.5); r2.italic = True
    r2.font.color.rgb = RGBColor(80, 80, 80)

    doc.save(OUTPUT_PATH)
    print(f"\nSaved: {OUTPUT_PATH}")
    import os
    sz = os.path.getsize(OUTPUT_PATH)
    print(f"Size: {sz:,} bytes ({sz/1024:.0f} KB)")


make_doc()
