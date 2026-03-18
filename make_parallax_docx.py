import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_PATH  = r"C:\Users\bryan\.claude\plans\polished-prancing-flame.md"
OUTPUT_PATH = r"C:\Users\bryan\Documents\Parallax_White_Paper_arXiv.docx"

PAPER_TITLE     = "Parallax: Community-Structured Graph Attention\nfor Knowledge Graph Reasoning"
AUTHOR_NAME     = "Bryan Alexander Buchorn"
AUTHOR_CALLSIGN = "AMP"
AFFILIATION     = "Independent Researcher"
EMAIL           = "bryan.alexander@buchorn.com"
COAUTHOR_NAME   = "Claude Sonnet 4.6 (Research Collaborator)"
COAUTHOR_AFF    = "Anthropic"
DATE_STR        = "March 2026"
ARXIV_NOTE      = "Preprint \u2014 Version 0.1"
COLOR_INLINE    = RGBColor(149, 55, 53)


def add_border_bottom(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom); pPr.append(pBdr)


def set_cell_shading(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)


def apply_inline(para, text, base_size=Pt(10)):
    for part in re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)").split(text):
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = para.add_run(part[2:-2]); r.bold = True; r.font.size = base_size
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = para.add_run(part[1:-1]); r.italic = True; r.font.size = base_size
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = para.add_run(part[1:-1]); r.font.name = "Courier New"
            r.font.size = Pt(9); r.font.color.rgb = COLOR_INLINE
        elif part:
            r = para.add_run(part); r.font.size = base_size


def setup_styles(doc):
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(10)
    for level, size, sb in [(1, 11, 10), (2, 10, 8), (3, 10, 6), (4, 10, 4)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"; h.font.size = Pt(size)
        h.font.bold = True; h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(sb); h.paragraph_format.space_after = Pt(3)
        h.paragraph_format.keep_with_next = True
    try:
        cs = doc.styles["Code Block"]
    except KeyError:
        cs = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    cs.font.name = "Courier New"; cs.font.size = Pt(8.5)
    cs.paragraph_format.left_indent = Inches(0.15)
    cs.paragraph_format.space_before = cs.paragraph_format.space_after = Pt(0)
    cs.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    cs.paragraph_format.line_spacing = Pt(11)


def enable_two_columns(doc):
    from docx.enum.section import WD_SECTION_START
    sec = doc.add_section(); sec.start_type = WD_SECTION_START.CONTINUOUS
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1.0)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "2"); cols.set(qn("w:space"), str(int(0.25 * 1440)))
    cols.set(qn("w:equalWidth"), "1"); sec._sectPr.append(cols)


def parse_table_rows(tlines):
    rows = []
    for line in tlines:
        if re.match(r"^\|[-: |]+\|$", line.strip()):
            continue
        rows.append([c.strip() for c in line.strip().strip("|").split("|")])
    return rows


def add_table(doc, rows):
    if not rows:
        return
    cc = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=cc); tbl.style = "Table Grid"
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            if ci >= cc:
                continue
            cell = tbl.rows[ri].cells[ci]; cell.text = ""; p = cell.paragraphs[0]; p.clear()
            apply_inline(p, ct, base_size=Pt(8.5))
            p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(1)
            if ri == 0:
                for run in p.runs: run.bold = True
                set_cell_shading(cell, "D3D3D3")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def make_doc():
    doc = Document(); setup_styles(doc)
    for sec in doc.sections:
        sec.page_width = Inches(8.5); sec.page_height = Inches(11)
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1.0)

    with open(INPUT_PATH, encoding="utf-8") as f:
        content = f.read()

    abm = re.search(r"## Abstract\s*\n(.*?)(?=\n---|\n## )", content, re.DOTALL)
    abstract_text = abm.group(1).strip() if abm else ""

    # Title
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_lines = PAPER_TITLE.split("\n")
    for j, tl in enumerate(title_lines):
        run = tp.add_run(tl); run.bold = True
        run.font.name = "Times New Roman"; run.font.size = Pt(16)
        if j < len(title_lines) - 1:
            tp.add_run("\n")
    tp.paragraph_format.space_before = Pt(0); tp.paragraph_format.space_after = Pt(8)

    # Primary author
    ap = doc.add_paragraph(); ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = ap.add_run(f"{AUTHOR_NAME}  \u00b7  {AUTHOR_CALLSIGN}")
    r1.font.name = "Times New Roman"; r1.font.size = Pt(11); r1.bold = True
    ap.add_run("     ")
    r2 = ap.add_run(COAUTHOR_NAME); r2.font.name = "Times New Roman"; r2.font.size = Pt(11)
    ap.paragraph_format.space_before = Pt(0); ap.paragraph_format.space_after = Pt(2)

    aff = doc.add_paragraph(); aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ra = aff.add_run(f"{AFFILIATION}          {COAUTHOR_AFF}")
    ra.font.name = "Times New Roman"; ra.font.size = Pt(10); ra.italic = True
    aff.paragraph_format.space_before = Pt(0); aff.paragraph_format.space_after = Pt(2)

    em = doc.add_paragraph(); em.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rem = em.add_run(EMAIL); rem.font.name = "Courier New"; rem.font.size = Pt(9)
    em.paragraph_format.space_before = Pt(0); em.paragraph_format.space_after = Pt(2)

    dv = doc.add_paragraph(); dv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rdv = dv.add_run(f"{DATE_STR}   \u00b7   {ARXIV_NOTE}")
    rdv.font.name = "Times New Roman"; rdv.font.size = Pt(9); rdv.italic = True
    dv.paragraph_format.space_before = Pt(0); dv.paragraph_format.space_after = Pt(8)
    add_border_bottom(dv)

    # Abstract
    abl = doc.add_paragraph(); abl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rabll = abl.add_run("Abstract"); rabll.bold = True
    rabll.font.name = "Times New Roman"; rabll.font.size = Pt(10)
    abl.paragraph_format.space_before = Pt(8); abl.paragraph_format.space_after = Pt(3)

    abp = doc.add_paragraph()
    abp.paragraph_format.left_indent = abp.paragraph_format.right_indent = Inches(0.5)
    abp.paragraph_format.space_before = Pt(0); abp.paragraph_format.space_after = Pt(12)
    apply_inline(abp, abstract_text, base_size=Pt(9.5))

    # Two-column body
    enable_two_columns(doc)

    lines = content.split("\n")
    i = 0; in_code = False; code_lines = []
    table_lines = []; in_table = False
    skip_abstract = False; skip_header = True

    while i < len(lines):
        line = lines[i]

        if skip_header:
            if re.match(r"^## Abstract", line):
                skip_header = False; skip_abstract = True
            i += 1; continue

        if skip_abstract:
            is_new_h2 = bool(re.match(r"^## ", line)) and "Abstract" not in line
            if line.strip() == "---" or is_new_h2:
                skip_abstract = False
                if not is_new_h2:
                    i += 1; continue
            else:
                i += 1; continue

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True; code_lines = []
            else:
                in_code = False
                for cl in code_lines:
                    p = doc.add_paragraph(style="Code Block"); p.add_run(cl if cl else " ")
                doc.add_paragraph().paragraph_format.space_after = Pt(4)
            i += 1; continue

        if in_code:
            code_lines.append(line.rstrip()); i += 1; continue

        if line.strip().startswith("|"):
            if not in_table:
                in_table = True; table_lines = []
            table_lines.append(line); i += 1
            if i < len(lines) and lines[i].strip().startswith("|"):
                continue
            in_table = False; add_table(doc, parse_table_rows(table_lines)); table_lines = []; continue

        if line.strip() == "---":
            i += 1; continue

        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1)); text = m.group(2)
            p = doc.add_heading(level=level); p.clear()
            apply_inline(p, text, base_size=Pt(11 if level == 1 else 10))
            i += 1; continue

        m = re.match(r"^([ \t]*)[-*]\s+(.*)", line)
        if m:
            ilvl = len(m.group(1)) // 2; text = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.15 + ilvl * 0.15)
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(1)
            p.clear(); apply_inline(p, text, base_size=Pt(10)); i += 1; continue

        m = re.match(r"^(\d+)\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(1)
            p.clear(); apply_inline(p, m.group(2), base_size=Pt(10)); i += 1; continue

        if line.strip() == "":
            i += 1; continue

        p = doc.add_paragraph()
        apply_inline(p, line.strip(), base_size=Pt(10))
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
        i += 1

    # References
    rh = doc.add_heading("References", level=1); rh.paragraph_format.space_before = Pt(10)
    refs = [
        '[1] Scarselli et al., "The Graph Neural Network Model," IEEE TNNLS, 2009.',
        '[2] Gilmer et al., "Neural Message Passing for Quantum Chemistry," ICML, 2017.',
        '[3] Velickovic et al., "Graph Attention Networks," ICLR, 2018.',
        '[4] Hamilton et al., "Inductive Representation Learning on Large Graphs," NeurIPS, 2017.',
        '[5] Bordes et al., "Translating Embeddings for Modeling Multi-relational Data (TransE)," NeurIPS, 2013.',
        '[6] Sun et al., "RotatE: Knowledge Graph Embedding by Relational Rotation in Complex Space," ICLR, 2019.',
        '[7] Xiong et al., "DeepPath: A Reinforcement Learning Method for Knowledge Graph Reasoning," EMNLP, 2017.',
        '[8] Das et al., "Go for a Walk and Arrive at the Answer (MINERVA)," ICLR, 2018.',
        '[9] Yao et al., "KG-GPT: A General Framework for Reasoning on Knowledge Graphs Using LLMs," 2023.',
        '[10] Chen et al., "KGPT: Knowledge-Grounded Pre-Training for Data-to-Text Generation," EMNLP, 2020.',
        '[11] Edge et al., "From Local to Global: A Graph RAG Approach to Query-Focused Summarization," Microsoft Research, 2024.',
        '[12] Sarthi et al., "RAPTOR: Recursive Abstractive Processing for Tree-Organized Retrieval," ICLR, 2024.',
        '[13] Blondel et al., "Fast Unfolding of Communities in Large Networks (Louvain)," JSTAT, 2008.',
        '[14] Traag et al., "From Louvain to Leiden: Guaranteeing Well-Connected Communities," Scientific Reports, 2019.',
        '[15] Raghavan et al., "Near Linear Time Algorithm to Detect Community Structures in Large-Scale Networks (LPA)," Physical Review E, 2007.',
        '[16] Galarraga et al., "AMIE: Association Rule Mining under Incomplete Evidence in Ontological Knowledge Bases," WWW, 2013.',
    ]
    for ref in refs:
        p = doc.add_paragraph(); r = p.add_run(ref)
        r.font.name = "Times New Roman"; r.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


make_doc()
