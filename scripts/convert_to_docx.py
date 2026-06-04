from typing import List
import sys
from docx import Document

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line:
                doc.add_paragraph(line)
    doc.save(docx_path)
    print(f"Created {docx_path}")

if __name__ == "__main__":
    if len(sys.argv) == 3:
        convert_md_to_docx(sys.argv[1], sys.argv[2])
    else:
        print("Usage: python convert_to_docx.py <md_path> <docx_path>")
