"""
Markdown â†’ .docx converter using python-docx.

Handles: headings, paragraphs, bold, italic, code blocks, tables,
horizontal rules, bullet lists, and inline LaTeX (rendered as italic text).

Usage:
    python scripts/md_to_docx.py docs/CEREBRUM_White_Paper.md
    # writes docs/CEREBRUM_White_Paper.docx
"""
from __future__ import annotations
from typing import List, Set

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def strip_latex(text: str) -> str:
    """Remove LaTeX delimiters, leaving the math expression as plain text."""
    # Block math: $$ ... $$
    text = re.sub(r'\$\$(.+?)\$\$', lambda m: '[' + m.group(1).strip() + ']', text, flags=re.DOTALL)
    # Inline math: $ ... $
    text = re.sub(r'\$(.+?)\$', lambda m: m.group(1).strip(), text)
    return text


def add_inline_run(para, text: str) -> None:
    """Parse basic inline markdown (bold, italic, code) and add runs."""
    # Split on ** bold **, * italic *, ` code `
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) >= 2:
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            if part:
                para.add_run(part)


def add_paragraph(doc: Document, text: str, style: str = 'Normal') -> None:
    text = strip_latex(text)
    # Remove markdown link syntax [text](url) â†’ text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    para = doc.add_paragraph(style=style)
    add_inline_run(para, text)


# ---------------------------------------------------------------------------
# Main converter
# ---------------------------------------------------------------------------

def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding='utf-8').splitlines()

    in_code_block = False
    code_lines: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        # Filter separator rows (---|---|...)
        data_rows = [r for r in table_rows if not re.match(r'^[\s\|\-:]+$', ''.join(r))]
        if not data_rows:
            in_table = False
            table_rows = []
            return
        cols = max(len(r) for r in data_rows)
        tbl = doc.add_table(rows=len(data_rows), cols=cols)
        tbl.style = 'Table Grid'
        for ri, row in enumerate(data_rows):
            for ci in range(cols):
                cell_text = row[ci].strip() if ci < len(row) else ''
                cell_text = strip_latex(cell_text)
                cell_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', cell_text)
                cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                cell_text = re.sub(r'\*(.+?)\*', r'\1', cell_text)
                cell_text = re.sub(r'`([^`]+)`', r'\1', cell_text)
                tbl.cell(ri, ci).text = cell_text.strip()
                if ri == 0:
                    for run in tbl.cell(ri, ci).paragraphs[0].runs:
                        run.bold = True
        doc.add_paragraph()  # spacing after table
        in_table = False
        table_rows = []

    def flush_code(lines_: list[str]) -> None:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        run = para.add_run('\n'.join(lines_))
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                flush_code(code_lines)
                code_lines = []
                in_code_block = False
            else:
                if in_table:
                    flush_table()
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table rows
        if line.strip().startswith('|'):
            in_table = True
            cells = [c for c in line.split('|')]
            # Remove leading/trailing empty cells from | delimiters
            if cells and not cells[0].strip():
                cells = cells[1:]
            if cells and not cells[-1].strip():
                cells = cells[:-1]
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # Horizontal rule
        if re.match(r'^---+\s*$', line):
            doc.add_paragraph('â”€' * 60)
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text = strip_latex(m.group(2))
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            f'Heading {min(level, 4)}'
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # Numbered section headers like "12.7 Bridge Twin Nodes"
        m = re.match(r'^(\d+\.[\d\.]*\s+\S.*)', line)
        if m and len(line) < 100 and not line.startswith(' '):
            text = strip_latex(line)
            para = doc.add_paragraph(style='Heading 2')
            add_inline_run(para, text)
            i += 1
            continue

        # Bullet lists
        m = re.match(r'^(\s*[-*+])\s+(.*)', line)
        if m:
            text = strip_latex(m.group(2))
            para = doc.add_paragraph(style='List Bullet')
            add_inline_run(para, text)
            i += 1
            continue

        # Numbered lists
        m = re.match(r'^\s*(\d+)[.)]\s+(.*)', line)
        if m:
            text = strip_latex(m.group(2))
            para = doc.add_paragraph(style='List Number')
            add_inline_run(para, text)
            i += 1
            continue

        # Empty line â†’ paragraph break
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Normal paragraph
        add_paragraph(doc, line)
        i += 1

    # Flush any remaining table/code
    if in_table:
        flush_table()
    if in_code_block and code_lines:
        flush_code(code_lines)

    doc.save(str(docx_path))
    print(f"  Written: {docx_path}")


if __name__ == '__main__':
    targets = [
        'docs/CEREBRUM_White_Paper.md',
        'docs/CEREBRUM_Whitepaper_V1.md',
        'docs/CEREBRUM_White_Paper_arXiv.md',
        'docs/CEREBRUM_Plain_Language_Guide.md',
    ]
    if len(sys.argv) > 1:
        targets = sys.argv[1:]

    for md_src in targets:
        src = Path(md_src)
        if not src.exists():
            print(f"  SKIP (not found): {src}")
            continue
        dst = src.with_suffix('.docx')
        print(f"Converting {src} -> {dst}")
        convert(src, dst)

    print("Done.")
